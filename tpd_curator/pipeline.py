import os
import re
import json
import yaml
from pathlib import Path
from multiprocessing import freeze_support
from datetime import datetime
from typing import Any, Dict, List, Set, Tuple, Union
import numpy as np
import pandas as pd
import time
import io
import csv
import pubchempy as pcp

from marker.config.parser import ConfigParser

from .preprocessing.convert_pdf_to_md import convert_pdf_to_md
from .preprocessing.data_cleaning import clean_json_data
from .llm.run_llm_api import *

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx is not available. DOCX file processing will be disabled.")

from langchain_core.messages import HumanMessage, AIMessage


def find_latest_cleaned_json_dir(output_dir: Path) -> str:
    """Find the latest timestamped cleaned_jsons directory"""
    cleaned_papers_dir = output_dir / "cleaned_papers"
    
    if not cleaned_papers_dir.exists():
        raise FileNotFoundError(f"Cleaned papers directory does not exist: {cleaned_papers_dir}")
    
    timestamp_dirs = []
    for item in cleaned_papers_dir.iterdir():
        if item.is_dir() and len(item.name) == 11 and item.name[6] == '_':
            try:
                datetime.strptime(item.name, "%y%m%d_%H%M")
                timestamp_dirs.append(item)
            except ValueError:
                continue
    
    if not timestamp_dirs:
        raise FileNotFoundError(f"No timestamped directories found in {cleaned_papers_dir}")
    
    latest_dir = sorted(timestamp_dirs, key=lambda d: d.name, reverse=True)[0]
    latest_cleaned_dir = latest_dir / "cleaned_jsons"
    
    if not latest_cleaned_dir.exists():
        raise FileNotFoundError(f"Cleaned JSONs directory does not exist: {latest_cleaned_dir}")
    
    print(f"Found latest cleaned JSONs directory: {latest_cleaned_dir}")
    return str(latest_cleaned_dir)



def load_config(config_path: str) -> dict:
    config_file = Path(config_path)
    if not config_file.exists():
        raise FileNotFoundError(f"Configuration file not found at {config_path}")
    with open(config_file, "r", encoding="utf-8") as file:
        config = yaml.safe_load(file)
    return config



def _read_text_head(file_path: Union[str, Path], n=5, encodings=('utf-8-sig','utf-8','latin1')) -> Tuple[str, str]:
    """Read first n lines of text file to determine delimiter/header, return (text,..."""
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc, errors='strict') as f:
                lines = ''.join([next(f) for _ in range(n)])
            return lines, enc
        except StopIteration:
            return lines, enc
        except Exception:
            continue
    with open(file_path, 'r', encoding='latin1', errors='ignore') as f:
        lines = ''.join([f.readline() for _ in range(n)])
    return lines, 'latin1'


def _sniff_delimiter(text_head: str) -> str:
    """Try to detect delimiter using csv"""
    try:
        dialect = csv.Sniffer().sniff(text_head, delimiters=[',',';','\t','|'])
        return dialect.delimiter
    except Exception:
        counts = {d: text_head.count(d) for d in [',',';','\t','|']}
        if counts[';'] >= max(counts.values()):
            return ';'
        if counts['\t'] >= max(counts.values()):
            return '\t'
        return ','


def _try_multiheader_read_csv(file_path, encoding, delimiter):
    """Try reading CSV with multi-level header [0,1]; fallback to single header if..."""
    df = pd.read_csv(
        file_path,
        sep=delimiter,
        encoding=encoding,
        engine='python',
        header=[0,1],
        dtype=str
    )
    if isinstance(df.columns, pd.MultiIndex):
        flatten_bad = sum(1 for lvl in df.columns.levels for v in lvl if isinstance(v, str) and v.startswith('Unnamed'))
        if flatten_bad > 0.3 * len(df.columns):
            df = pd.read_csv(
                file_path,
                sep=delimiter,
                encoding=encoding,
                engine='python',
                header=0,
                dtype=str
            )
    return df



def _flatten_columns(cols) -> list:
    """Flatten multi-level columns: ('BRD4','pDC50') -> 'BRD4 | pDC50'"""
    if isinstance(cols, pd.MultiIndex):
        flat = []
        for tup in cols:
            parts = [str(x).strip() for x in tup if pd.notna(x) and str(x).strip() and not str(x).startswith('Unnamed')]
            name = ' | '.join(parts) if parts else 'unnamed'
            flat.append(name)
        return flat
    else:
        return [str(c).strip() for c in cols]


def _coerce_numeric(df: pd.DataFrame) -> pd.DataFrame:
    """Convert numeric-looking columns to numeric, handling comma as decimal point..."""
    out = df.copy()
    for c in out.columns:
        s = out[c]
        sample = ''.join((s.dropna().astype(str).head(20)).tolist())
        letters = sum(ch.isalpha() for ch in sample)
        digits  = sum(ch.isdigit() for ch in sample)
        if digits == 0 or letters > digits:
            continue
        s2 = s.astype(str).str.replace(r'\s+', '', regex=True)
        s2 = s2.str.replace('%', '', regex=False)
        american_thousand = s2.str.contains(r'\d,\d{3}(?!\d)', regex=True, na=False).any()
        if american_thousand:
            s2 = s2.str.replace(r',(\d{3})(?!\d)', r'\1', regex=True)
        elif s2.str.count(',').sum() > s2.str.count(r'\.').sum():
            s2 = s2.str.replace('.', '', regex=False)
            s2 = s2.str.replace(',', '.', regex=False)
        out[c] = pd.to_numeric(s2, errors='ignore')
    return out


def read_markdown_file(file_path: Union[str, Path]) -> str:
    """Read markdown file and return its content as string"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content
    except Exception as e:
        print(f"Warning: Failed to read markdown file {file_path}: {e}")
        return ""


def read_csv_excel_file(file_path: Union[str, Path]) -> str:
    """Read CSV/Excel file and return formatted text representation"""
    file_path = Path(file_path)
    file_ext = file_path.suffix.lower()

    try:
        if file_ext == '.csv':
            head_text, enc = _read_text_head(file_path)
            delim = _sniff_delimiter(head_text)
            df = _try_multiheader_read_csv(file_path, enc, delim)

            df.columns = _flatten_columns(df.columns)
            df = df.loc[:, ~df.columns.duplicated()]
            df = df.dropna(axis=1, how='all').dropna(axis=0, how='all')

            df = _coerce_numeric(df)

            output = f"\n=== Table from {file_path.name} ===\n"
            output += f"Shape: {df.shape[0]} rows × {df.shape[1]} columns\n"
            output += f"Columns: {', '.join(df.columns.tolist())}\n\n"
            output += df.to_string(index=False, max_rows=100)
            output += "\n"

            return output

        elif file_ext in ['.xlsx', '.xls']:
            excel_file = pd.ExcelFile(file_path)
            sheet_names = excel_file.sheet_names

            output = f"\n=== Excel file: {file_path.name} ===\n"
            output += f"Total sheets: {len(sheet_names)}\n"

            sheets_content = []

            for sheet_name in sheet_names:
                try:
                    try:
                        df = pd.read_excel(file_path, sheet_name=sheet_name, header=[0,1], dtype=str)
                        if isinstance(df.columns, pd.MultiIndex):
                            flatten_bad = sum(1 for lvl in df.columns.levels for v in lvl if isinstance(v, str) and v.startswith('Unnamed'))
                            if flatten_bad > 0.3 * len(df.columns):
                                df = pd.read_excel(file_path, sheet_name=sheet_name, header=0, dtype=str)
                    except Exception:
                        df = pd.read_excel(file_path, sheet_name=sheet_name, header=0, dtype=str)

                    df.columns = _flatten_columns(df.columns)
                    df = df.loc[:, ~df.columns.duplicated()]
                    df = df.dropna(axis=1, how='all').dropna(axis=0, how='all')

                    df = _coerce_numeric(df)

                    sheet_content = f"--- Sheet: {sheet_name} ---\n"
                    sheet_content += f"Shape: {df.shape[0]} rows × {df.shape[1]} columns\n"
                    sheet_content += f"Columns: {', '.join(df.columns.tolist())}\n\n"
                    sheet_content += df.to_string(index=False, max_rows=100)
                    sheet_content += "\n\n"

                    sheets_content.append({
                        'name': sheet_name,
                        'content': sheet_content
                    })

                except Exception as e:
                    sheet_content = f"--- Sheet: {sheet_name} ---\n"
                    sheet_content += f"Error reading sheet: {str(e)}\n\n"
                    sheets_content.append({
                        'name': sheet_name,
                        'content': sheet_content
                    })

            for sheet_info in sheets_content:
                output += sheet_info['content']

            return output

        else:
            return f"Unsupported file format: {file_ext}"

    except Exception as e:
        print(f"Warning: Failed to read table file {file_path}: {e}")
        return f"Error reading {file_path.name}: {str(e)}\n"


def read_docx_file(file_path: Union[str, Path]) -> str:
    """Read DOCX file and return text content including paragraphs and tables"""
    if not DOCX_AVAILABLE:
        return f"Cannot read DOCX file {file_path}: python-docx not available\n"

    try:
        doc = Document(file_path)
        output = f"\n=== Document from {Path(file_path).name} ===\n\n"

        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        if paragraphs:
            output += "Text content:\n"
            output += "\n".join(paragraphs)
            output += "\n\n"

        if doc.tables:
            output += f"Tables found: {len(doc.tables)}\n\n"
            for i, table in enumerate(doc.tables, 1):
                output += f"Table {i}:\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    output += row_text + "\n"
                output += "\n"

        return output

    except Exception as e:
        print(f"Warning: Failed to read DOCX file {file_path}: {e}")
        return f"Error reading {Path(file_path).name}: {str(e)}\n"


def _contains_relevant_keywords(content: str) -> bool:
    """Check if content contains any of the relevant keywords: DC50, Dmax"""
    if not content:
        return False

    keywords_pattern = r'\b(DC50|Dmax)\b'

    match = re.search(keywords_pattern, content, re.IGNORECASE)

    return match is not None


def _filter_excel_sheets_by_keywords(file_path: Union[str, Path]) -> str:
    """Read Excel file and return only sheets containing relevant keywords"""
    try:
        excel_file = pd.ExcelFile(file_path)
        sheet_names = excel_file.sheet_names

        relevant_sheets = []
        filtered_sheets = []

        for sheet_name in sheet_names:
            try:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name, header=[0,1], dtype=str)
                    if isinstance(df.columns, pd.MultiIndex):
                        flatten_bad = sum(1 for lvl in df.columns.levels for v in lvl if isinstance(v, str) and v.startswith('Unnamed'))
                        if flatten_bad > 0.3 * len(df.columns):
                            df = pd.read_excel(file_path, sheet_name=sheet_name, header=0, dtype=str)
                except Exception:
                    df = pd.read_excel(file_path, sheet_name=sheet_name, header=0, dtype=str)

                df.columns = _flatten_columns(df.columns)
                df = df.loc[:, ~df.columns.duplicated()]
                df = df.dropna(axis=1, how='all').dropna(axis=0, how='all')

                df = _coerce_numeric(df)

                sheet_content = f"--- Sheet: {sheet_name} ---\n"
                sheet_content += f"Shape: {df.shape[0]} rows × {df.shape[1]} columns\n"
                sheet_content += f"Columns: {', '.join(df.columns.tolist())}\n\n"
                sheet_content += df.to_string(index=False, max_rows=100)
                sheet_content += "\n\n"

                if _contains_relevant_keywords(sheet_content):
                    relevant_sheets.append({
                        'name': sheet_name,
                        'content': sheet_content
                    })
                else:
                    filtered_sheets.append(sheet_name)

            except Exception as e:
                print(f"    Warning: Failed to read sheet '{sheet_name}': {e}")

        if not relevant_sheets:
            if filtered_sheets:
                print(f"    All {len(filtered_sheets)} sheets filtered out (no keywords)")
            return ""

        output = f"\n=== Excel file: {Path(file_path).name} ===\n"
        output += f"Total sheets: {len(sheet_names)}\n"
        output += f"Relevant sheets: {len(relevant_sheets)}\n"
        if filtered_sheets:
            output += f"Filtered out: {len(filtered_sheets)} sheets\n"
        output += "\n"

        for sheet_info in relevant_sheets:
            output += sheet_info['content']
            print(f"    ✓ Including sheet '{sheet_info['name']}' (contains keywords)")

        for sheet_name in filtered_sheets:
            print(f"    ✗ Filtering out sheet '{sheet_name}' (no keywords)")

        return output

    except Exception as e:
        print(f"    Error filtering Excel sheets: {e}")
        return ""


def load_supplementary_files_for_pmc(pmc_id: str, json_path: str, filter_by_keywords: bool = True) -> List[Dict[str, str]]:
    """Load all supplementary files for a given PMC ID and return list of..."""
    supplementary_folder = Path(json_path) / pmc_id / "supplementary"

    if not supplementary_folder.exists():
        print(f"No supplementary folder found for {pmc_id}")
        return []

    supplementary_files = []
    files_filtered_out = 0

    try:
        for file_path in supplementary_folder.iterdir():
            if file_path.name == "supplementary_metadata.json":
                continue

            if not file_path.is_file():
                continue

            file_ext = file_path.suffix.lower()
            content = ""

            if file_ext == '.md':
                content = read_markdown_file(file_path)

                if content:
                    supplementary_files.append({
                        'filename': file_path.name,
                        'content': f"\n--- Markdown: {file_path.name} ---\n{content}\n",
                        'file_type': 'markdown'
                    })
                    print(f"  ✓ Including {file_path.name} (markdown - for IUPAC name extraction)")

            elif file_ext in ['.xlsx', '.xls']:
                if filter_by_keywords:
                    print(f"  Processing {file_path.name} (filtering by sheet)...")
                    content = _filter_excel_sheets_by_keywords(file_path)
                    if content:
                        supplementary_files.append({
                            'filename': file_path.name,
                            'content': content,
                            'file_type': 'excel'
                        })
                    else:
                        files_filtered_out += 1
                else:
                    content = read_csv_excel_file(file_path)
                    if content:
                        supplementary_files.append({
                            'filename': file_path.name,
                            'content': content,
                            'file_type': 'excel'
                        })

            elif file_ext == '.csv':
                content = read_csv_excel_file(file_path)

                if content:
                    if filter_by_keywords:
                        if _contains_relevant_keywords(content):
                            supplementary_files.append({
                                'filename': file_path.name,
                                'content': content,
                                'file_type': 'csv'
                            })
                            print(f"  ✓ Including {file_path.name} (contains DC50/Dmax)")
                        else:
                            files_filtered_out += 1
                            print(f"  ✗ Filtering out {file_path.name} (no relevant keywords found)")
                    else:
                        supplementary_files.append({
                            'filename': file_path.name,
                            'content': content,
                            'file_type': 'csv'
                        })

            elif file_ext == '.docx':
                content = read_docx_file(file_path)

                if content:
                    supplementary_files.append({
                        'filename': file_path.name,
                        'content': content,
                        'file_type': 'docx'
                    })
                    print(f"  ✓ Including {file_path.name} (docx - for IUPAC name extraction)")
            else:
                continue

        if not supplementary_files:
            if files_filtered_out > 0:
                print(f"No relevant supplementary files found for {pmc_id} (filtered out {files_filtered_out} files/sheets without keywords)")
            else:
                print(f"No processable supplementary files found for {pmc_id}")
            return []

        print(f"Loaded {len(supplementary_files)} supplementary file(s) for {pmc_id}")
        if filter_by_keywords and files_filtered_out > 0:
            print(f"Filtered out {files_filtered_out} files/sheets without relevant keywords")

        return supplementary_files

    except Exception as e:
        print(f"Error loading supplementary files for {pmc_id}: {e}")
        return []



def initialize_character_tracking(pmc_id: str, with_history: bool) -> dict:
    """Initialize character tracking data structure for a paper"""
    return {
        "pmc_id": pmc_id,
        "with_history": with_history,
        "rounds": [],
        "summary": {
            "total_rounds_counted": 0,
            "total_characters_all_rounds": 0,
            "total_prompt_characters": 0,
            "total_paper_characters": 0
        }
    }


def record_llm_input_characters(tracking_data: dict,
                                round_number: int,
                                round_type: str,
                                description: str,
                                complete_prompt: str,
                                custom_prompt: str = None,
                                paper_text: str = None,
                                supplementary_filename: str = None) -> None:
    """Record character counts for a single LLM input round"""
    total_chars = len(complete_prompt) if complete_prompt else 0
    prompt_chars = len(custom_prompt) if custom_prompt else 0
    paper_chars = len(paper_text) if paper_text else 0

    round_data = {
        "round_number": round_number,
        "round_type": round_type,
        "description": description,
        "supplementary_filename": supplementary_filename,
        "total_characters": total_chars,
        "prompt_characters": prompt_chars,
        "paper_characters": paper_chars
    }

    tracking_data["rounds"].append(round_data)

    summary = tracking_data["summary"]
    summary["total_rounds_counted"] += 1
    summary["total_characters_all_rounds"] += total_chars
    summary["total_prompt_characters"] += prompt_chars
    summary["total_paper_characters"] += paper_chars

    print(f"  Character stats - Total: {total_chars}, Prompt: {prompt_chars}, Paper: {paper_chars}")


def save_character_length_json(tracking_data: dict, output_path: Path) -> None:
    """Save character tracking data to JSON file"""
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(tracking_data, f, ensure_ascii=False, indent=2)
        print(f"\n{'='*60}")
        print(f"Character Length Statistics Summary:")
        print(f"  Total rounds: {tracking_data['summary']['total_rounds_counted']}")
        print(f"  Total characters: {tracking_data['summary']['total_characters_all_rounds']:,}")
        print(f"  Total prompt characters: {tracking_data['summary']['total_prompt_characters']:,}")
        print(f"  Total paper characters: {tracking_data['summary']['total_paper_characters']:,}")
        print(f"Saved to: {output_path}")
        print(f"{'='*60}\n")
    except Exception as e:
        print(f"Error saving character length JSON to {output_path}: {e}")




class InMemoryHistory(BaseChatMessageHistory, BaseModel):
    messages: list[BaseMessage] = Field(default_factory=list)
    def add_messages(self, messages: list[BaseMessage]) -> None:
        self.messages.extend(messages)
    def clear(self) -> None:
        self.messages = []

store = {}

def get_by_session_id(session_id: str) -> BaseChatMessageHistory:
    if session_id not in store:
        store[session_id] = InMemoryHistory()
    return store[session_id]


def test_history_memory(session_id: str):
    print(f"\n--- Verify session history for: {session_id} ---")

    history = get_by_session_id(session_id)

    if not isinstance(history, InMemoryHistory):
        print("❌ get_session_history did not return an InMemoryHistory instance")
        return

    if not history.messages:
        print("⚠️ Current session has no history messages.")
        return

    print(f"✅ Current session has {len(history.messages)} messages:")
    for i, msg in enumerate(history.messages):
        role = "👤User" if isinstance(msg, HumanMessage) else "🤖AI"
        print(f"  {i+1:02d}. {role}: {msg.content[:80]}{'...' if len(msg.content) > 80 else ''}")


def clean_json_papers_batch(config: dict):
    """Process JSON papers with data cleaning"""
    json_path = config["data"].get("json_path")
    if not json_path or not Path(json_path).exists():
        raise FileNotFoundError(f"JSON data path does not exist or not specified: {json_path}")
    
    print(f"Starting batch processing of JSON papers from: {json_path}")
    
    output_dir = Path(config["data"].get("output_dir", "./output"))
    
    result = clean_json_data(
        input_path=json_path,
        output_dir=output_dir,
        fields_to_remove=config["data"].get("json_fields_to_remove"),
        sections_to_remove=config["data"].get("sections_to_remove"),
        figure_fields_to_remove=config["data"].get("figure_fields_to_remove"),
        paper_filter_file=config["data"].get("paper_filter_file")
    )
    
    cleaned_data, cleaned_output_dir = result
    
    if not cleaned_data:
        print("No JSON data was successfully cleaned.")
        return 0
    
    print(f"Completed batch processing of {len(cleaned_data)} JSON papers.")
    print(f"Files saved to: {cleaned_output_dir}")
    
    config["data"]["clean_json_dir"] = cleaned_output_dir
    print(f"Updated clean_json_dir to: {cleaned_output_dir}")
    
    print(f"{'-'*80}")
    
    return len(cleaned_data)


def run_llm_pipeline(config: dict, read_papers: bool = False, md_prompt: str = None, with_history: bool = False, tutorial_text: str = None, no_timestamp: bool = False, custom_output_dir: Path = None):
    """Run LLM to extract molecular glue data"""

    if no_timestamp and custom_output_dir:
        results_dir = Path(custom_output_dir)
        os.makedirs(results_dir, exist_ok=True)
    else:
        timestamp = datetime.now().strftime("%y%m%d_%H%M")
        results_dir = Path(config["results"]["results_dir"]) / timestamp
        os.makedirs(results_dir, exist_ok=True)
    config_yaml_path = results_dir / "config.yaml"

    try:
        with open(config_yaml_path, "w", encoding="utf-8") as f:
            yaml.dump(config, f, allow_unicode=True, sort_keys=False)
        print(f"Saved config to {config_yaml_path}")
    except Exception as e:
        print(f"Failed to save config YAML: {e}")

    if read_papers:
        if tutorial_text:
            print(f"\nUsing tutorial with {len(tutorial_text)} characters")
            prompt_text = tutorial_text + "\n\n---\n\n" + md_prompt
        else:
            prompt_text = md_prompt
    
        json_file_dir = Path(config["data"]["clean_json_dir"])
        
        if not json_file_dir.exists() or not json_file_dir.is_dir():
            print(f"Configured clean_json_dir does not exist: {json_file_dir}")
            try:
                output_dir = Path(config["data"].get("output_dir", "./output"))
                latest_dir = find_latest_cleaned_json_dir(output_dir)
                json_file_dir = Path(latest_dir)
                print(f"Using latest cleaned JSONs directory: {json_file_dir}")
            except Exception as e:
                print(f"Failed to find latest cleaned directory: {e}")
                return

        json_files = sorted(json_file_dir.glob("*.json"))
        if not json_files:
            print(f"No JSON files found in {json_file_dir}")
            return
        print(f"Running LLM reasoning on {len(json_files)} files in {json_file_dir} ...")

        for json_file in json_files:
            with open(json_file, "r", encoding="utf-8") as f:
                data = json.load(f)
            records, model_name = None, "" 
            filename = json_file.stem
            if filename.endswith("_cleaned"):
                file_number = filename.replace("_cleaned", "")
            else:
                file_number = filename

            file_paper_dir = results_dir / file_number
            file_paper_dir.mkdir(parents=True, exist_ok=True)

            tracking_data = initialize_character_tracking(
                pmc_id=file_number,
                with_history=with_history if with_history else False
            )

            file_content = json.dumps(data, ensure_ascii=False, indent=2)
            
            if not file_content.strip():
                print(f"Warning: No content found in {json_file.name}")
                continue      

            print(f"\n{'='*60}")
            print(f"Run LLM reasoning on file: {json_file.name}")
            print(f"File content length: {len(file_content)} characters")
            
            print("Using LLM API for reasoning...")
            if with_history:
                print("Using LLM API with history...")

                session_id = f"session_{file_number}"

                first_round_prompt = construct_prompt(
                    custom_prompt=prompt_text,
                    paper_text=file_content
                )

                record_llm_input_characters(
                    tracking_data=tracking_data,
                    round_number=1,
                    round_type="main_text_extraction",
                    description="Structured data extraction from main text",
                    complete_prompt=first_round_prompt,
                    custom_prompt=prompt_text,
                    paper_text=file_content,
                    supplementary_filename=None
                )

                print(f"\n{'='*60} INPUT TO LLM {'='*60}")
                print(first_round_prompt)
                print(f"{'='*130}")

                records1, model_name = run_llm_api_with_history(
                    session_id=session_id,
                    custom_prompt=first_round_prompt,
                    get_session_history=get_by_session_id,
                    temperature=0
                )

                print(f"\n{'='*60} LLM OUTPUT {'='*60}")
                print(records1)
                print(f"{'='*130}")

                if records1:
                    out1 = file_paper_dir / f"{file_number}_extraction.json"
                    with out1.open("w", encoding="utf-8") as f:
                        json.dump(records1, f, ensure_ascii=False, indent=2)
                    print(f"Saved extraction → {out1}")

                second_round_prompt = construct_prompt(
                    custom_prompt=(
                        "For each extracted data point, find the exact source sentence (verbatim) where you found that data. "
                        "Return JSON ONLY as an array of objects with keys \"data_point\" and \"source_sentence\". "
                        "Do not paraphrase the sentence.\n"
                        "Example: "
                        "[{\"data_point\":{"
                        "\"Compound_Name\":\"...\","
                        "\"IUPAC_Name\":\"...\","
                        "\"SMILES\":\"...\","
                        "\"Degradation_Target\":\"...\","
                        "\"Recruiter\":\"...\","
                        "\"Assay\":\"...\","
                        "\"Cell_Line\":\"...\","
                        "\"DC50\":\"...\","
                        "\"DC50_units\":\"...\","
                        "\"DC50_h\":\"...\","
                        "\"Dmax\":\"...\","
                        "\"Dmax_h\":\"...\","
                        "\"Dmax_conc\":\"...\""
                        "},"
                        "\"source_sentence\":\"...\"}]"
                    )
                    )


                records2, _ = run_llm_api_with_history(
                    session_id=session_id,
                    custom_prompt=second_round_prompt,
                    get_session_history=get_by_session_id,
                    temperature=0
                )
                test_history_memory(session_id=session_id)


                if records2:
                    if isinstance(records2, str):
                        out2 = file_paper_dir / f"{file_number}_source_sentence.md"
                        out2.write_text(records2, encoding="utf-8")
                    else:
                        out2 = file_paper_dir / f"{file_number}_source_sentence.json"
                        with out2.open("w", encoding="utf-8") as f:
                            json.dump(records2, f, ensure_ascii=False, indent=2)
                    print(f"Saved source sentences → {out2}")

                if config["data"].get("include_supplementary", False):
                    print(f"\n{'='*60} ROUND 3+: SUPPLEMENTARY FILES {'='*60}")

                    json_path = config["data"]["json_path"]
                    filter_by_keywords = config["data"].get("filter_supplementary_by_keywords", True)

                    if filter_by_keywords:
                        print("Filtering supplementary files by keywords (DC50, Dmax)...")
                    else:
                        print("Loading all supplementary files without keyword filtering...")

                    supplementary_files = load_supplementary_files_for_pmc(
                        pmc_id=file_number,
                        json_path=json_path,
                        filter_by_keywords=filter_by_keywords
                    )

                    if supplementary_files:
                        print(f"\nFound {len(supplementary_files)} supplementary file(s) to process")
                        print("Processing each file sequentially...\n")

                        all_supplementary_results = []

                        for idx, supp_file in enumerate(supplementary_files, 1):
                            filename = supp_file['filename']
                            content = supp_file['content']
                            file_type = supp_file.get('file_type', 'unknown')

                            print(f"\n{'='*60} ROUND {2 + idx}: Supplementary File {idx}/{len(supplementary_files)} {'='*60}")
                            print(f"Processing: {filename} (type: {file_type})")

                            if file_type in ['markdown', 'docx']:
                                file_type_display = "markdown" if file_type == "markdown" else "Word"
                                supp_custom_prompt = f"""Review your previous extraction and check this supplementary {file_type_display} file #{idx} ({filename}):

**Primary Task: Find IUPAC Names and SMILES for Previously Extracted Compounds**

Please carefully examine this supplementary file and:
1. Look for IUPAC names, SMILES strings, or full chemical names for compounds you previously extracted (e.g., if you extracted "compound 5", look for its IUPAC name, SMILES string, or chemical structure name)
2. For each compound where you find an IUPAC name or SMILES:
   - Keep all other fields (Compound_Name, Degradation_Target, DC50, Dmax, etc.) EXACTLY the same as your previous extraction
   - ONLY fill in or update the IUPAC_Name and SMILES fields
   - This helps match compounds across different parts of the paper

**Secondary Task: Extract Any New Compound Data**
3. If you find any NEW molecular glue degradation data that was NOT in your previous extraction, extract it completely
4. Include any additional compounds, assays, or cell lines found in this file

**Important Notes:**
- When updating existing compounds with IUPAC names or SMILES, preserve all original data
- If a compound already has an IUPAC name or SMILES, only update if you find a more complete or accurate one
- Return all data points (both updated and new) in the same JSON format as before
- Output a single JSON array of objects, with no added commentary or explanation

If no additional data or IUPAC names are found in this file, return an empty array: []"""
                            else:
                                supp_custom_prompt = f"""Review your previous extraction and check this supplementary file #{idx} ({filename}):

Question: Are there any data points you missed in the first extraction?

Please carefully examine this supplementary file and:
1. Identify any molecular glue degradation data that was NOT in your previous extraction
2. Extract any missing fields (DC50, Dmax, Cell_Line, Assay, etc.) that can now be filled
3. Look for IUPAC names, SMILES strings, or full chemical names for compounds you previously extracted
4. Include any new compounds, assays, or cell lines found in this file

**Important for IUPAC Names and SMILES:**
- If you find IUPAC names or SMILES strings for compounds you already extracted, fill in the IUPAC_Name and SMILES fields while keeping other fields the same

Return all data points (both updated and new) in the same JSON format as before. Output a single JSON array of objects, with no added commentary or explanation

If no additional data is found in this file, return an empty array: []"""

                            third_round_prompt = construct_prompt(
                                custom_prompt=supp_custom_prompt,
                                paper_text=content
                            )

                            record_llm_input_characters(
                                tracking_data=tracking_data,
                                round_number=2 + idx,
                                round_type="supplementary_extraction",
                                description=f"Supplementary file #{idx}",
                                complete_prompt=third_round_prompt,
                                custom_prompt=supp_custom_prompt,
                                paper_text=content,
                                supplementary_filename=filename
                            )

                            print(f"\n{'='*60} INPUT TO LLM (Round {2 + idx}) {'='*60}")
                            print(f"File: {filename}")
                            print(f"Content length: {len(content)} characters")
                            print(f"{'='*130}")

                            try:
                                records_supp, _ = run_llm_api_with_history(
                                    session_id=session_id,
                                    custom_prompt=third_round_prompt,
                                    get_session_history=get_by_session_id,
                                    temperature=0
                                )

                                print(f"\n{'='*60} LLM OUTPUT (Round {2 + idx}) {'='*60}")
                                print(records_supp)
                                print(f"{'='*130}")

                                if records_supp:
                                    if isinstance(records_supp, str):
                                        print(f"Warning: LLM returned text instead of structured JSON for {filename}")
                                        print(f"This may indicate the file contains no relevant data or LLM failed to parse.")

                                        text_out = file_paper_dir / f"{file_number}_supplementary_{idx}_{filename}_text_response.txt"
                                        with text_out.open("w", encoding="utf-8") as f:
                                            f.write(records_supp)
                                        print(f"Saved text response to: {text_out}")
                                    else:
                                        all_supplementary_results.append({
                                            'filename': filename,
                                            'file_index': idx,
                                            'extraction': records_supp
                                        })

                            except Exception as e:
                                print(f"Error processing {filename}: {e}")
                                print(f"Continuing with next file...")
                                continue

                        if all_supplementary_results:
                            out3 = file_paper_dir / f"{file_number}_supplementary_extraction.json"
                            with out3.open("w", encoding="utf-8") as f:
                                json.dump(all_supplementary_results, f, ensure_ascii=False, indent=2)
                            print(f"\nSaved all supplementary extractions → {out3}")
                            print(f"Total supplementary files processed: {len(all_supplementary_results)}")

                        test_history_memory(session_id=session_id)

                        print(f"\n{'='*60} FINAL REVIEW ROUND {'='*60}")
                        print("Reviewing main text to find missing IUPAC names and SMILES for extracted compounds...")

                        round_number = 2 + len(supplementary_files) + 1

                        final_review_custom_prompt = """Review all the compounds you've extracted so far from both the main text and supplementary files.

**Task: Find Missing IUPAC Names and SMILES in Main Text**

Please carefully examine the main text below and:
1. Identify any compounds you extracted that are still missing IUPAC names or SMILES strings
2. Search the main text for IUPAC names, SMILES strings, or full chemical names for these compounds
3. For each compound where you find an IUPAC name or SMILES:
   - Keep all other fields (Compound_Name, Degradation_Target, DC50, Dmax, etc.) EXACTLY the same as your previous extraction
   - ONLY fill in or update the IUPAC_Name and SMILES fields
   - Match the compound by its name (e.g., "compound 5", "molecule A") or by its degradation data

**Important:**
- Focus on compounds extracted from supplementary files that might have IUPAC names or SMILES in the main text
- Preserve all existing data for each compound
- Only update the IUPAC_Name and SMILES fields; do not modify any other fields
- Return all data points (both updated and unchanged) in the same JSON format

If no additional IUPAC names or SMILES are found, return an empty array: []"""

                        final_review_prompt = construct_prompt(
                            custom_prompt=final_review_custom_prompt,
                            paper_text=file_content
                        )

                        record_llm_input_characters(
                            tracking_data=tracking_data,
                            round_number=round_number,
                            round_type="final_review",
                            description="Final review to find missing IUPAC names and SMILES in main text",
                            complete_prompt=final_review_prompt,
                            custom_prompt=final_review_custom_prompt,
                            paper_text=file_content,
                            supplementary_filename=None
                        )

                        print(f"\n{'='*60} INPUT TO LLM (Final Review) {'='*60}")
                        print(final_review_prompt)
                        print(f"{'='*130}")

                        try:
                            records_final_review, _ = run_llm_api_with_history(
                                session_id=session_id,
                                custom_prompt=final_review_prompt,
                                get_session_history=get_by_session_id,
                                temperature=0
                            )

                            print(f"\n{'='*60} LLM OUTPUT (Final Review) {'='*60}")
                            print(records_final_review)
                            print(f"{'='*130}")

                            out_final = file_paper_dir / f"{file_number}_final_review.json"
                            if isinstance(records_final_review, list):
                                with out_final.open("w", encoding="utf-8") as f:
                                    json.dump(records_final_review, f, ensure_ascii=False, indent=2)
                                print(f"\nSaved final review extraction → {out_final}")
                                if records_final_review:
                                    print(f"Final review found {len(records_final_review)} compounds with updated IUPAC names/SMILES")
                                else:
                                    print("No additional IUPAC names or SMILES found in main text")
                            else:
                                with out_final.open("w", encoding="utf-8") as f:
                                    f.write(str(records_final_review))
                                print(f"\nSaved final review result (non-list format) → {out_final}")

                        except Exception as e:
                            print(f"Error in Final Review Round: {e}")
                            import traceback
                            traceback.print_exc()

                        print(f"{'='*60} END FINAL REVIEW ROUND {'='*60}\n")

                    else:
                        print(f"No supplementary files found for {file_number}, skipping supplementary rounds.")
                else:
                    print("Supplementary file processing is disabled in config.")

                char_stats_path = file_paper_dir / "character_length.json"
                save_character_length_json(tracking_data, char_stats_path)

            else:
                print("Using LLM API without history...")

                complete_prompt = construct_prompt(
                    custom_prompt=prompt_text,
                    paper_text=file_content
                )

                record_llm_input_characters(
                    tracking_data=tracking_data,
                    round_number=1,
                    round_type="main_text_extraction",
                    description="Single round extraction (no history)",
                    complete_prompt=complete_prompt,
                    custom_prompt=prompt_text,
                    paper_text=file_content,
                    supplementary_filename=None
                )

                print(f"\n{'='*60} INPUT TO LLM {'='*60}")
                print(f"Custom prompt: {prompt_text}")
                print(f"\nPaper text: {file_content}")
                print(f"{'='*130}")

                records, model_name = run_llm_api(
                custom_prompt = prompt_text,
                paper_text = file_content)
                
                print(f"\n{'='*60} LLM OUTPUT {'='*60}")
                print(records)
                print(f"{'='*130}")

                if records:
                    for r in records:
                        r["file_number"] = file_number
                        r["model_name"] = model_name

                    output_json_file = file_paper_dir / f"{file_number}_answer.json"
                    with open(output_json_file, "w", encoding="utf-8") as out_f:
                        json.dump(records, out_f, ensure_ascii=False, indent=2)
                    print(f"Saved {len(records)} record(s) to {output_json_file}")

                char_stats_path = file_paper_dir / "character_length.json"
                save_character_length_json(tracking_data, char_stats_path)

    else:
        prompt_text = config["prompt"]["default_prompt"]

        llm_output = run_llm_api(custom_prompt=prompt_text)
        if llm_output:
            output_json_file = results_dir / f"answer.json"
            try:
                with open(output_json_file, "w", encoding="utf-8") as out_f:
                    json.dump(llm_output, out_f, ensure_ascii=False, indent=2)
                print(f"LLM output saved to: {output_json_file}")
            except Exception as e:
                print(f"Failed to save JSON: {e}")

    print("-" * 60)
    print("All reasoning tasks finished.")
    print("-" * 60)


def save_csv(config_file: dict, results_dir: str = None):
    """Generate CSV files from JSON extraction results"""
    if results_dir:
        latest_json_dir = Path(results_dir)
        if not latest_json_dir.exists():
            raise ValueError(f"Specified results directory does not exist: {results_dir}")
        print(f"Loading JSON files from the specified directory: {latest_json_dir}")
    else:
        json_file_path = Path(config_file["results"]["results_dir"])
        if not json_file_path.exists():
            raise ValueError(f"Results directory does not exist: {json_file_path}")

        all_subdirs = [d for d in json_file_path.iterdir() if d.is_dir()]
        if not all_subdirs:
            raise ValueError(f"No result subdirectories found under: {json_file_path}")

        timestamped_dirs = []
        for subdir in all_subdirs:
            match = re.search(r"(\d{6})_(\d{4})", subdir.name)
            if match:
                timestamp_value = int(match.group(1) + match.group(2))
                timestamped_dirs.append((timestamp_value, subdir))

        if timestamped_dirs:
            timestamped_dirs.sort(key=lambda item: item[0], reverse=True)
            latest_json_dir = timestamped_dirs[0][1]
            print(f"Loading JSON files from the latest directory (by timestamp): {latest_json_dir}")
        else:
            latest_json_dir = max(all_subdirs, key=lambda d: d.stat().st_mtime)
            print(f"Loading JSON files from the latest directory (by modification time): {latest_json_dir}")

    csv_data_file = latest_json_dir / "data.csv"

    pmc_to_doi_map = {}
    paper_ids_file = config_file.get('data', {}).get('paper_ids_file', None)
    if paper_ids_file:
        molglue_papers_path = Path(paper_ids_file)
    else:
        molglue_papers_path = Path("/Users/yaochenr/project/tpd_curator/data_source/250825_molglue_papers/paper_dois/avaliable_paper_ids.json")

    if molglue_papers_path.exists():
        try:
            with open(molglue_papers_path, 'r') as f:
                papers_data = json.load(f)
                for paper in papers_data:
                    if 'pmc_id' in paper and 'doi' in paper:
                        pmc_to_doi_map[paper['pmc_id']] = paper['doi']
            print(f"Loaded {len(pmc_to_doi_map)} PMC ID to DOI mappings")
        except Exception as e:
            print(f"Failed to load PMC to DOI mapping: {e}")

    data_list = []

    for json_file in latest_json_dir.rglob("*.json"):
        if "source_sentence" in json_file.name:
            print(f"Skipping source sentence file: {json_file.name}")
            continue

        if "character_length" in json_file.name:
            print(f"Skipping character length metadata file: {json_file.name}")
            continue

        if not (json_file.name.endswith("_extraction.json") or
                json_file.name.endswith("_supplementary_extraction.json") or
                json_file.name.endswith("_final_review.json")):
            print(f"Skipping non-extraction file: {json_file.name}")
            continue

        try:
            with open(json_file, "r", encoding="utf-8") as f:
                content = json.load(f)

                filename = json_file.stem

                if filename.endswith("_supplementary_extraction"):
                    pmc_id = filename.replace("_supplementary_extraction", "")
                    doi = pmc_to_doi_map.get(pmc_id, "Unknown")

                    if isinstance(content, list):
                        for supp_file in content:
                            if isinstance(supp_file, dict) and 'extraction' in supp_file:
                                extraction = supp_file['extraction']

                                if isinstance(extraction, list):
                                    for record in extraction:
                                        if isinstance(record, dict):
                                            record['DOI'] = doi
                                            record['source'] = f"supplementary_{supp_file.get('filename', 'unknown')}"
                                    data_list.extend(extraction)
                                elif isinstance(extraction, dict):
                                    extraction['DOI'] = doi
                                    extraction['source'] = f"supplementary_{supp_file.get('filename', 'unknown')}"
                                    data_list.append(extraction)

                    print(f"Merged supplementary data from: {json_file.name}")
                    continue

                if filename.endswith("_final_review"):
                    pmc_id = filename.replace("_final_review", "")
                    doi = pmc_to_doi_map.get(pmc_id, "Unknown")

                    if isinstance(content, list):
                        for record in content:
                            if isinstance(record, dict):
                                record['DOI'] = doi
                                record['source'] = 'final_review'
                        data_list.extend(content)
                    elif isinstance(content, dict):
                        content['DOI'] = doi
                        content['source'] = 'final_review'
                        data_list.append(content)

                    print(f"Merged final review data from: {json_file.name}")
                    continue

                if filename.endswith("_extraction"):
                    pmc_id = filename.replace("_extraction", "")
                else:
                    pmc_id = filename

                doi = pmc_to_doi_map.get(pmc_id, "Unknown")

                if isinstance(content, list):
                    for record in content:
                        if isinstance(record, dict):
                            record['DOI'] = doi
                            record['source'] = 'main_text'
                    data_list.extend(content)
                elif isinstance(content, dict):
                    content['DOI'] = doi
                    content['source'] = 'main_text'
                    data_list.append(content)
                else:
                    print(f"Skip {json_file}, it's content is not a list or dict.")
                    continue

        except Exception as e:
            print(f"Failed to read or parse {json_file}: {e}")

    if not data_list:
        print("Cannot find any valid JSON data in the specified directory.")
        return

    df = pd.DataFrame(data_list)

    priority_columns = []
    if 'DOI' in df.columns:
        priority_columns.append('DOI')
    if 'source' in df.columns:
        priority_columns.append('source')

    if priority_columns:
        other_columns = [col for col in df.columns if col not in priority_columns]
        df = df[priority_columns + other_columns]
        print(f"Added {', '.join(priority_columns)} column(s) at the beginning")

        if 'source' in df.columns:
            source_counts = df['source'].value_counts()
            print(f"\nData source statistics:")
            for source, count in source_counts.items():
                print(f"  {source}: {count} records")

    df.to_csv(csv_data_file, index=False, encoding="utf-8")
    print(f"\nSaved all LLM output (main text + supplementary) to: {Path(csv_data_file)}")
    print(f"Total records: {len(df)}")

    return csv_data_file


_APPROX_PREFIX_RE = re.compile(
    r'^\s*(?:about|approximately|approx\.?|around|circa|ca\.|nearly|near)\s*(?=[\d.])',
    re.IGNORECASE,
)


def convert_to_nm(value_str: str, unit: str) -> str:
    """Convert concentration value to nM, handling comparison operators,..."""
    try:
        unit = normalize_concentration_unit(unit)
        if unit is None:
            return value_str

        if isinstance(value_str, str):
            value_str = _APPROX_PREFIX_RE.sub('~', value_str.strip())

        if unit == 'nm':
            operator_match = re.match(r'^([><=≥≤≈~∼]+\s*)(.*)', value_str.strip())
            if operator_match:
                operator = operator_match.group(1).strip() + ' '
                number_part = operator_match.group(2).strip()
            else:
                operator = ''
                number_part = value_str.strip()

            number_part = re.sub(r'[a-zA-Zµμ%]+', '', number_part).strip()
            number_part = number_part.replace(',', '')

            return f"{operator}{number_part}".strip()

        operator_match = re.match(r'^([><=≥≤≈~∼]+\s*)(.*)', value_str.strip())
        if operator_match:
            operator = operator_match.group(1).strip() + ' '
            number_part = operator_match.group(2).strip()
        else:
            operator = ''
            number_part = value_str.strip()

        number_part = re.sub(r'[a-zA-Zµμ%]+', '', number_part).strip()

        number_part = number_part.replace(',', '')

        def apply_conversion(val):
            """Apply unit conversion factor to a numeric value (unit is already normalized)"""
            if unit == 'um':
                return val * 1000
            elif unit == 'mm':
                return val * 1000000
            elif unit == 'm':
                return val * 1000000000
            elif unit == 'pm':
                return val / 1000
            elif unit == 'fm':
                return val / 1000000
            else:
                return None

        def format_value(val):
            """Format converted value preserving original precision"""
            result = str(val)
            if '.' in result:
                result = result.rstrip('0').rstrip('.')
            return result

        if '±' in number_part:
            parts = number_part.split('±')
            if len(parts) == 2:
                main_value = float(parts[0].strip())
                uncertainty = float(parts[1].strip())

                converted_main = apply_conversion(main_value)
                converted_unc = apply_conversion(uncertainty)

                if converted_main is None:
                    return value_str

                if converted_main >= 1000:
                    return f"{operator}{converted_main:.0f} ± {converted_unc:.0f}"
                elif converted_main >= 100:
                    return f"{operator}{converted_main:.1f} ± {converted_unc:.1f}"
                else:
                    return f"{operator}{converted_main:.2f} ± {converted_unc:.2f}"

        elif re.search(r'(?<!^)(?<![\s])-(?!$)', number_part):
            parts = re.split(r'\s*-\s*', number_part)
            if len(parts) == 2:
                try:
                    low = float(parts[0].strip())
                    high = float(parts[1].strip())

                    converted_low = apply_conversion(low)
                    converted_high = apply_conversion(high)

                    if converted_low is None:
                        return value_str

                    return f"{operator}{format_value(converted_low)}-{format_value(converted_high)}"
                except ValueError:
                    return value_str

        else:
            value = float(number_part)
            converted_value = apply_conversion(value)

            if converted_value is None:
                return value_str

            return f"{operator}{format_value(converted_value)}"

    except Exception:
        return value_str

_DMAX_OPERATOR_CHARS = '~><≥≤≈∼='
_DMAX_OPERATOR_RE = re.compile(f'[{_DMAX_OPERATOR_CHARS}]')

_DMAX_NULL_RE = re.compile(r'^\s*no[\s\-]*degradation\b', re.IGNORECASE)
_DMAX_PHRASE_PATTERNS = [
    (re.compile(r'^\s*up[\s\-]*to\s*(?=[\d.])', re.IGNORECASE), '<='),
    (re.compile(r'^\s*over\s*(?=[\d.])', re.IGNORECASE), '>'),
]


def clean_dmax_value(value_str):
    """Normalize Dmax while preserving operators"""
    if pd.isna(value_str) or value_str == 'null':
        return None

    s = str(value_str).strip()

    if _DMAX_NULL_RE.match(s):
        return None

    for pattern, replacement in _DMAX_PHRASE_PATTERNS:
        s = pattern.sub(replacement, s)

    s = _APPROX_PREFIX_RE.sub('~', s)

    s = re.sub(
        r'(?:between\s+)?(\d+(?:\.\d+)?)\s*(?:and|to)\s*(\d+(?:\.\d+)?)',
        r'\1-\2',
        s,
        flags=re.IGNORECASE,
    )
    s = s.replace('–', '-').replace('—', '-')

    s = s.replace('≥', '>=').replace('≤', '<=').replace('∼', '~')

    s = re.sub(r'[%\s]', '', s)

    return s if s else None


def normalize_dmax_for_dedup(val):
    """Strip operator characters from Dmax for dedup-key construction only"""
    if pd.isna(val):
        return ''
    return _DMAX_OPERATOR_RE.sub('', str(val).strip())


def dmax_has_operator(val):
    """True if the Dmax string contains any comparison/approximation operator"""
    if pd.isna(val):
        return False
    return bool(_DMAX_OPERATOR_RE.search(str(val)))


def is_pdc50_unit(unit_str) -> bool:
    """True if the unit string denotes pDC50 / log10[M] (i"""
    if unit_str is None or (isinstance(unit_str, float) and pd.isna(unit_str)):
        return False
    u = str(unit_str).lower()
    return (
        'pdc50' in u
        or 'log10[m]' in u
        or 'log10(m)' in u
        or 'log10 m' in u
        or 'logm' in u
    )


def convert_pdc50_to_nm(value_str: str) -> str:
    """Convert a pDC50 (log10[M]) value to DC50 in nM"""
    if value_str is None or (isinstance(value_str, float) and pd.isna(value_str)) or value_str == 'null':
        return value_str

    s = str(value_str).strip()
    s = re.sub(r'^\s*p?dc\s*50\s*', '', s, flags=re.IGNORECASE).strip()
    s = re.sub(r'^\s*50\s+', '', s).strip()

    def to_nm(pval: float) -> float:
        return 10 ** (9 - pval)

    def fmt(v: float) -> str:
        return f"{v:.1f}"

    try:
        if '±' in s:
            center = float(s.split('±')[0].strip())
            return fmt(to_nm(center))
        if re.search(r'\d\s*-\s*\d', s):
            parts = re.split(r'\s*-\s*', s)
            if len(parts) == 2:
                low_p = float(parts[0].strip())
                high_p = float(parts[1].strip())
                return f"{fmt(to_nm(high_p))}-{fmt(to_nm(low_p))}"
        return fmt(to_nm(float(s)))
    except (ValueError, TypeError):
        return value_str


def clean_dc50_value(value_str: str) -> str:
    """Remove embedded units and clean formatting from DC50 values"""
    if pd.isna(value_str) or value_str == 'null':
        return value_str

    value_str = str(value_str).strip()

    operator_match = re.match(r'^([><=≥≤≈~∼]+\s*)(.*)', value_str)
    if operator_match:
        operator = operator_match.group(1).strip()
        operator = operator.replace('≥', '>=').replace('≤', '<=').replace('∼', '~')
        operator = operator + ' '
        number_part = operator_match.group(2).strip()
    else:
        operator = ''
        number_part = value_str

    number_part = re.sub(
        r'(\d+(?:\.\d+)?)\s*(?:and|to)\s*(\d+(?:\.\d+)?)',
        r'\1-\2',
        number_part,
        flags=re.IGNORECASE,
    )
    number_part = number_part.replace('–', '-').replace('—', '-')

    number_part = re.sub(r'[a-zA-Zµμ%]+', '', number_part).strip()

    number_part = number_part.replace(',', '')

    return f"{operator}{number_part}".strip()

def process_units(csv_file_path: str):
    print("-" * 60)
    try:
        df = pd.read_csv(csv_file_path)

        if 'DC50' in df.columns and 'DC50_units' not in df.columns:
            print("DC50_units column not found. Extracting units from DC50...")
            temp_dc50_value = []
            temp_dc50_units = []

            for idx, row in df.iterrows():
                if pd.notna(row['DC50']) and row['DC50'] != 'null':
                    dc50_str = str(row['DC50'])

                    match = re.search(r'(.*?)([a-zA-Z\s%µμ]+)$', dc50_str.strip())
                    if match:
                        value_str = match.group(1).strip()
                        unit = match.group(2).strip()

                        if unit.lower() != 'nm':
                            converted_value = convert_to_nm(value_str, unit)
                            if converted_value != value_str:
                                temp_dc50_value.append(converted_value)
                                temp_dc50_units.append('nM')
                            else:
                                temp_dc50_value.append(value_str)
                                temp_dc50_units.append(unit)
                        else:
                            temp_dc50_value.append(value_str)
                            temp_dc50_units.append('nM')
                    else:
                        temp_dc50_value.append(dc50_str.strip())
                        temp_dc50_units.append(None)
                else:
                    temp_dc50_value.append(None)
                    temp_dc50_units.append(None)

            dc50_idx = df.columns.get_loc('DC50')

            df.drop(columns=['DC50'], inplace=True)
            df.insert(dc50_idx, 'DC50', temp_dc50_value)
            df.insert(dc50_idx + 1, 'DC50_units', temp_dc50_units)
        elif 'DC50' in df.columns and 'DC50_units' in df.columns:
            print("DC50_units column already exists. Processing all DC50 values...")

            for idx, row in df.iterrows():
                if pd.notna(row['DC50_units']) and row['DC50_units'] != 'null':
                    unit = str(row['DC50_units']).strip()

                    if pd.notna(row['DC50']) and row['DC50'] != 'null':
                        dc50_value = str(row['DC50']).strip()

                        if is_pdc50_unit(unit) or is_pdc50_unit(dc50_value):
                            converted_value = convert_pdc50_to_nm(dc50_value)
                            if converted_value != dc50_value:
                                df.at[idx, 'DC50'] = converted_value
                                df.at[idx, 'DC50_units'] = 'nM'
                                print(f"  Converted: {dc50_value} {unit} → {converted_value} nM (pDC50)")
                            else:
                                print(f"  Warning: Could not convert pDC50 value {dc50_value}")
                            continue

                        has_embedded_unit = bool(re.search(r'[a-zA-Zµμ%]', dc50_value))

                        if has_embedded_unit or unit.lower() not in ['nm', 'nanomolar']:

                            if unit.lower() in ['nm', 'nanomolar']:
                                cleaned_value = clean_dc50_value(dc50_value)
                                if cleaned_value != dc50_value:
                                    df.at[idx, 'DC50'] = cleaned_value
                                    print(f"  Cleaned: {dc50_value} → {cleaned_value} (unit: nM)")
                            else:
                                converted_value = convert_to_nm(dc50_value, unit)

                                if converted_value != dc50_value:
                                    df.at[idx, 'DC50'] = converted_value
                                    df.at[idx, 'DC50_units'] = 'nM'
                                    print(f"  Converted: {dc50_value} {unit} → {converted_value} nM")
                                else:
                                    cleaned_value = clean_dc50_value(dc50_value)
                                    if cleaned_value != dc50_value:
                                        df.at[idx, 'DC50'] = cleaned_value
                                        print(f"  Warning: Could not convert unit {unit}, but cleaned: {dc50_value} → {cleaned_value}")
                                    else:
                                        print(f"  Warning: Could not convert {dc50_value} {unit} to nM")

            print("DC50 processing complete.")

        if 'Dmax' in df.columns:
            temp_dmax_value = []

            for idx, row in df.iterrows():
                if pd.notna(row['Dmax']) and row['Dmax'] != 'null':
                    temp_dmax_value.append(clean_dmax_value(row['Dmax']))
                else:
                    temp_dmax_value.append(None)

            dmax_idx = df.columns.get_loc('Dmax')

            df.drop(columns=['Dmax'], inplace=True)
            df.insert(dmax_idx, 'Dmax', temp_dmax_value)

        if 'Dmax_concentration' in df.columns:
            df.rename(columns={'Dmax_concentration': 'Dmax_conc'}, inplace=True)

        print(f"\nDeduplicating records...")
        before_count = len(df)
        print(f"Before deduplication: {before_count} records")

        potential_key_fields = ['Compound_Name', 'Degradation_Target', 'Recruiter', 'Assay', 'Cell_Line', 'DC50', 'Dmax']
        key_fields = [field for field in potential_key_fields if field in df.columns]

        if len(key_fields) < 2:
            print(f"Warning: Not enough key fields for deduplication (found: {key_fields}). Skipping deduplication.")
            key_fields = None
        else:
            print(f"Using key fields for deduplication: {key_fields}")

        _bookkeeping_cols = {'_dedup_key', '_dedup_base', '_dedup_info', '_sub_key', 'source'}

        def _row_dedup_key(row):
            parts = []
            for f in key_fields:
                val = row[f] if pd.notna(row[f]) else ''
                if f == 'Dmax':
                    parts.append(normalize_dmax_for_dedup(val))
                else:
                    parts.append(str(val))
            return '|'.join(parts)

        if key_fields:
            df['_dedup_key'] = df.apply(_row_dedup_key, axis=1)

        def _non_empty_field_count(row):
            count = 0
            for col in row.index:
                if col in _bookkeeping_cols:
                    continue
                val = row[col]
                if pd.isna(val):
                    continue
                s = str(val).strip()
                if s == '' or s.lower() == 'null':
                    continue
                count += 1
            return count

        def _compute_base_rank(row):
            has_iupac = (
                'IUPAC_Name' in row.index
                and pd.notna(row['IUPAC_Name'])
                and row['IUPAC_Name'] != ''
            )
            source = str(row.get('source', '')) if 'source' in row.index else ''
            if has_iupac:
                if source == 'final_review':
                    return 0
                if 'supplementary' in source.lower():
                    return 1
                return 2
            if source == 'final_review':
                return 3
            if 'supplementary' in source.lower():
                return 4
            return 5

        if key_fields:
            df['_dedup_base'] = df.apply(_compute_base_rank, axis=1)
            df['_dedup_info'] = df.apply(_non_empty_field_count, axis=1)

            winners = []
            for _, group in df.groupby('_dedup_key', sort=False):
                group_sorted = group.sort_values(
                    ['_dedup_base', '_dedup_info'], ascending=[True, False]
                )
                winner = group_sorted.iloc[0].copy()

                if 'Dmax' in winner.index and dmax_has_operator(winner['Dmax']):
                    winner_bare = normalize_dmax_for_dedup(winner['Dmax'])
                    for _, other in group_sorted.iloc[1:].iterrows():
                        other_dmax = other.get('Dmax')
                        if pd.isna(other_dmax):
                            continue
                        if dmax_has_operator(other_dmax):
                            continue
                        if normalize_dmax_for_dedup(other_dmax) == winner_bare:
                            winner['Dmax'] = other_dmax
                            break

                winners.append(winner)

            df = pd.DataFrame(winners).reset_index(drop=True)
            df = df.drop(columns=['_dedup_key', '_dedup_base', '_dedup_info'])

            print(f"After deduplication: {len(df)} records")
            print(f"Removed {before_count - len(df)} duplicate records\n")

        identity_fields = ['Compound_Name', 'Degradation_Target', 'Recruiter', 'Assay', 'Cell_Line']
        identity_fields = [f for f in identity_fields if f in df.columns]

        value_fields = ['DC50', 'DC50_units', 'DC50_h', 'Dmax', 'Dmax_h', 'Dmax_conc', 'EC50', 'EC50_units', 'EC50_h']
        value_fields = [f for f in value_fields if f in df.columns]

        if len(identity_fields) >= 2 and len(value_fields) >= 1:
            print(f"Running subsumption deduplication...")
            print(f"  Identity fields: {identity_fields}")
            print(f"  Value fields: {value_fields}")
            before_sub = len(df)

            def _is_empty(val):
                """Check if a value is considered empty"""
                if val is None:
                    return True
                if isinstance(val, float) and pd.isna(val):
                    return True
                if isinstance(val, str) and val.strip() == '':
                    return True
                return False

            df['_sub_key'] = df[identity_fields].fillna('').astype(str).agg('|'.join, axis=1)
            rows_to_drop = set()

            for _, group_indices in df.groupby('_sub_key').groups.items():
                if len(group_indices) < 2:
                    continue
                indices = list(group_indices)
                row_info = {}
                for idx in indices:
                    non_empty = {}
                    for vf in value_fields:
                        val = df.at[idx, vf]
                        if not _is_empty(val):
                            if vf == 'Dmax':
                                non_empty[vf] = normalize_dmax_for_dedup(val)
                            else:
                                non_empty[vf] = str(val).strip()
                    row_info[idx] = non_empty

                for i in range(len(indices)):
                    idx_i = indices[i]
                    if idx_i in rows_to_drop:
                        continue
                    for j in range(len(indices)):
                        if i == j:
                            continue
                        idx_j = indices[j]
                        if idx_j in rows_to_drop:
                            continue
                        fields_i = row_info[idx_i]
                        fields_j = row_info[idx_j]
                        if set(fields_j.keys()).issubset(set(fields_i.keys())):
                            shared_match = all(
                                fields_i[k] == fields_j[k] for k in fields_j
                            )
                            if shared_match and len(fields_i) > len(fields_j):
                                rows_to_drop.add(idx_j)

            df = df.drop(index=rows_to_drop).reset_index(drop=True)
            df = df.drop(columns=['_sub_key'])
            print(f"  After subsumption dedup: {len(df)} records (removed {before_sub - len(df)})\n")
        else:
            print("Skipping subsumption deduplication due to insufficient fields.\n")

        processed_file_path = str(csv_file_path).replace('.csv', '_processed.csv')
        df.to_csv(processed_file_path, index=False, encoding="utf-8")
        print(f"Finish processing units, save results to {processed_file_path}")

        return processed_file_path

    except Exception as e:
        print(f"Error processing units:{e}")
        return csv_file_path


def normalize_unicode_chars(text: str) -> str:
    """Normalize Unicode characters that might appear identical but have different..."""
    if pd.isna(text):
        return text

    text = str(text).strip()
    text = text.replace('μ', 'µ')
    text = text.replace('µ', 'µ')

    import re
    text = re.sub(r'\s+', '', text)

    return text


def normalize_dashes(text: str) -> str:
    """Normalize all dash/hyphen variants to standard hyphen-minus (U+002D)"""
    if pd.isna(text):
        return text
    text = str(text)
    text = re.sub(r'[\u2010\u2011\u2012\u2013\u2014\u2015\u2212\uFE58\uFE63\uFF0D]', '-', text)
    return text



def extract_value_with_operator(text: str) -> tuple:
    """Extract comparison operator and numerical value from a string"""
    import re
    if pd.isna(text):
        return None, None

    text = str(text).strip()
    match = re.match(r'^([<>≤≥][=]?|[~∼≈])?\s*(\d+(?:\.\d+)?)', text)
    if match:
        operator = match.group(1) or ''
        number = float(match.group(2))

        operator_normalization = {
            '≥': '>=',
            '≤': '<=',
            '~': '',
            '∼': '',
            '≈': '',
        }
        operator = operator_normalization.get(operator, operator)

        return operator, number
    return None, None



def compare_exact_fields_with_normalization(val1, val2) -> bool:
    """Compare two field values with Unicode normalization and numerical value..."""
    if pd.isna(val1) and pd.isna(val2):
        return True
    elif pd.isna(val1) or pd.isna(val2):
        return False
    else:
        norm_val1 = normalize_unicode_chars(val1)
        norm_val2 = normalize_unicode_chars(val2)

        if norm_val1.lower() == norm_val2.lower():
            return True

        op1, num1 = extract_value_with_operator(norm_val1)
        op2, num2 = extract_value_with_operator(norm_val2)

        if num1 is not None and num2 is not None:
            numbers_match = abs(num1 - num2) < 1e-6

            if numbers_match:
                if op1 == op2:
                    return True
                elif op1 == '' or op2 == '':
                    return True
                else:
                    return False

        return norm_val1.lower() == norm_val2.lower()




def load_semantic_cache(cache_file_path: str = None) -> Dict[str, List[Dict]]:
    """Load the semantic-match cache file, grouped by DOI"""
    if cache_file_path is None:
        print("Warning: No cache file path provided, using empty cache")
        return {}

    try:
        if os.path.exists(cache_file_path):
            with open(cache_file_path, 'r', encoding='utf-8') as f:
                cache_data = json.load(f)
                if isinstance(cache_data, dict):
                    return cache_data
                elif isinstance(cache_data, list):
                    print(f"Warning: Cache file {cache_file_path} is in legacy format, migrating to DOI-grouped format")
                    return {"_legacy": cache_data}
                else:
                    print(f"Warning: Cache file {cache_file_path} format is invalid, starting with empty cache")
                    return {}
        else:
            print(f"Cache file {cache_file_path} does not exist, starting with empty cache")
            return {}
    except Exception as e:
        print(f"Error loading cache file {cache_file_path}: {e}, starting with empty cache")
        return {}


def save_semantic_cache(cache_data: Dict[str, List[Dict]], cache_file_path: str = None):
    """Save the semantic-match cache to disk, grouped by DOI"""
    if cache_file_path is None:
        print("Warning: No cache file path provided, cannot save cache")
        return

    try:
        with open(cache_file_path, 'w', encoding='utf-8') as f:
            json.dump(cache_data, f, ensure_ascii=False, indent=2)
        print(f"Cache saved to {cache_file_path}")
    except Exception as e:
        print(f"Error saving cache file {cache_file_path}: {e}")


def find_cache_match(pred_value: str, labeled_value: str, feature: str, cache_data: Dict[str, List[Dict]], doi: str = None) -> str:
    """Look up a match in the cache"""
    pred_value = str(pred_value).strip()
    labeled_value = str(labeled_value).strip()

    def search_in_list(items: List[Dict]) -> str:
        """Search a list of cache entries for a match"""
        for item in items:
            cache_pred = str(item['pred_value']).strip()
            cache_labeled = str(item['labeled_value']).strip()

            if (cache_pred.lower() == pred_value.lower() and
                cache_labeled.lower() == labeled_value.lower() and
                item.get('feature', '') == feature):
                return item['result']

            if (cache_pred.lower() == labeled_value.lower() and
                cache_labeled.lower() == pred_value.lower() and
                item.get('feature', '') == feature):
                return item['result']
        return None

    if doi and doi in cache_data:
        result = search_in_list(cache_data[doi])
        if result is not None:
            return result

    for doi_key, items in cache_data.items():
        if doi_key == doi:
            continue
        result = search_in_list(items)
        if result is not None:
            return result

    return None

def check_semantic_match(pred_value: str, labeled_value: str, feature: str = None, llm_api_key: str = None, llm_api_url: str = None, config: dict = None, cache_file_path: str = None, doi: str = None) -> bool:
    """Check whether two values are semantically equivalent, with a file cache to..."""

    if pd.isna(pred_value) or pd.isna(labeled_value):
        return False

    pred_value = str(pred_value).strip()
    labeled_value = str(labeled_value).strip()

    if pred_value.lower() == labeled_value.lower():
        return True

    if not feature:
        feature = "compound, protein or cell line"

    cache_data = load_semantic_cache(cache_file_path)

    cached_result = find_cache_match(pred_value, labeled_value, feature, cache_data, doi=doi)
    if cached_result is not None:
        print(f"Cache hit! The comparison result for {pred_value} and {labeled_value} ({feature}) is: {cached_result}")
        return cached_result == 'yes'

    prompt = f"""
    Please determine whether {pred_value} and {labeled_value} refer to the same {feature}.
    Please only answer "yes" or "no". If they are referring to the same thing, respond 'yes'; otherwise, respond 'no', don't add any extra words.
    """

    for _ in range(3):
        try:
            result, _ = run_llm_api(custom_prompt=prompt, paper_text=None)

            if not result:
                print("Response from LLM API is empty, retrying...")
                time.sleep(10)
                result, _ = run_llm_api(custom_prompt=prompt, paper_text=None)

            if result and isinstance(result, str):
                is_match = "yes" in result.lower()
                result_text = 'yes' if is_match else 'no'

                cache_item = {
                    'pred_value': pred_value,
                    'labeled_value': labeled_value,
                    'feature': feature,
                    'result': result_text
                }

                cache_key = doi if doi else "_global"
                if cache_key not in cache_data:
                    cache_data[cache_key] = []
                cache_data[cache_key].append(cache_item)
                save_semantic_cache(cache_data, cache_file_path)
                print(f"Added to cache ({cache_key}): The comparison result for {pred_value} and {labeled_value} ({feature}) is {result_text}")

                return is_match
            else:
                print(f"Cannot parse LLM's result: {result}, continue with exact matching.")
                return pred_value.lower() == labeled_value.lower()

        except Exception as e:
            print(f"Cannot use LLM's API, continue with exact matching:{e}")
            return pred_value.lower() == labeled_value.lower()


def is_missing_value(value) -> bool:
    """Check whether a value is empty/missing"""
    if pd.isna(value):
        return True
    if isinstance(value, str):
        return value.strip() == "" or value.lower() in ['null', 'none', 'nan', 'n/a']
    return False




def parse_value_with_range(value: str) -> Tuple[float, float, float]:
    """Parse a numeric string, supporting several interval formats"""
    if pd.isna(value) or is_missing_value(value):
        return (None, None, None)

    value_str = str(value).strip()

    match = re.match(r'^(\d+(?:\.\d+)?)\s*[±+]/?[-/]?\s*(\d+(?:\.\d+)?)', value_str)
    if match:
        center = float(match.group(1))
        error = float(match.group(2))
        return (center, center - error, center + error)

    match = re.match(r'^[\[\(](\d+(?:\.\d+)?)\s*,\s*(\d+(?:\.\d+)?)[\]\)]', value_str)
    if match:
        lower = float(match.group(1))
        upper = float(match.group(2))
        center = (lower + upper) / 2
        return (center, lower, upper)

    match = re.match(r'^(\d+(?:\.\d+)?)\s*[-~]\s*(\d+(?:\.\d+)?)', value_str)
    if match:
        lower = float(match.group(1))
        upper = float(match.group(2))
        center = (lower + upper) / 2
        return (center, lower, upper)

    match = re.match(r'^(\d+(?:\.\d+)?)', value_str)
    if match:
        center = float(match.group(1))
        return (center, center, center)

    return (None, None, None)


def check_value_in_range(pred_value: str, ref_value: str) -> bool:
    """Check whether pred_value falls inside ref_value's interval (or vice versa)"""
    pred_center, pred_lower, pred_upper = parse_value_with_range(pred_value)
    ref_center, ref_lower, ref_upper = parse_value_with_range(ref_value)

    if pred_center is None or ref_center is None:
        return False

    if ref_lower <= pred_center <= ref_upper:
        return True

    if pred_lower <= ref_center <= pred_upper:
        return True

    return False


def compare_assay_values(pred_value, ref_value) -> bool:
    """Compare assay values using exact match plus interval match"""
    if pd.isna(pred_value) and pd.isna(ref_value):
        return True
    elif pd.isna(pred_value) or pd.isna(ref_value):
        return False

    if compare_exact_fields_with_normalization(pred_value, ref_value):
        return True

    return check_value_in_range(str(pred_value), str(ref_value))


_CONCENTRATION_UNIT_ALIASES = {
    'nm': 'nm', 'nanomolar': 'nm',
    'nmol/l': 'nm', 'nmol.l-1': 'nm', 'nmoll-1': 'nm', 'nmol/liter': 'nm',
    'um': 'um', 'micromolar': 'um',
    'umol/l': 'um', 'umol.l-1': 'um', 'umoll-1': 'um', 'umol/liter': 'um',
    'mm': 'mm', 'millimolar': 'mm',
    'mmol/l': 'mm', 'mmol.l-1': 'mm', 'mmoll-1': 'mm', 'mmol/liter': 'mm',
    'm': 'm', 'molar': 'm',
    'mol/l': 'm', 'mol.l-1': 'm', 'moll-1': 'm', 'mol/liter': 'm',
    'pm': 'pm', 'picomolar': 'pm',
    'pmol/l': 'pm', 'pmol.l-1': 'pm', 'pmoll-1': 'pm', 'pmol/liter': 'pm',
    'fm': 'fm', 'femtomolar': 'fm',
    'fmol/l': 'fm', 'fmol.l-1': 'fm', 'fmoll-1': 'fm', 'fmol/liter': 'fm',
}


def normalize_concentration_unit(unit) -> str:
    """Normalize a concentration unit to a short token (e"""
    if unit is None or (isinstance(unit, float) and pd.isna(unit)):
        return None
    s = str(unit).strip().lower()
    if not s:
        return None
    s = s.replace('µ', 'u').replace('μ', 'u')
    s = re.sub(r'\s+', '', s)
    return _CONCENTRATION_UNIT_ALIASES.get(s, s)


def compare_concentration_units(pred_value, ref_value) -> bool:
    """Concentration-unit equivalence: 'nmol/L' == 'nM', 'μM' == 'umol/L', etc"""
    if pd.isna(pred_value) and pd.isna(ref_value):
        return True
    if pd.isna(pred_value) or pd.isna(ref_value):
        return False
    pn = normalize_concentration_unit(pred_value)
    rn = normalize_concentration_unit(ref_value)
    if pn is None and rn is None:
        return True
    if pn is None or rn is None:
        return False
    return pn == rn





def calculate_assay_match_score(pred_row, ref_row, assay_fields: list = None,
                                 tiebreak_fields: list = None) -> int:
    """Score the assay-value field match between two records"""
    if assay_fields is None:
        assay_fields = ['DC50', 'Dmax', 'Dmax_h', 'Dmax_conc']

    match_count = 0

    for field in assay_fields:
        if field not in pred_row or field not in ref_row:
            continue

        pred_val = pred_row[field]
        ref_val = ref_row[field]

        if is_missing_value(pred_val) and is_missing_value(ref_val):
            continue

        if is_missing_value(pred_val) or is_missing_value(ref_val):
            continue

        if compare_assay_values(pred_val, ref_val):
            match_count += 1

    if tiebreak_fields:
        for field in tiebreak_fields:
            if field not in pred_row or field not in ref_row:
                continue
            pred_val = pred_row[field]
            ref_val = ref_row[field]
            if is_missing_value(pred_val) or is_missing_value(ref_val):
                continue
            if compare_exact_fields_with_normalization(pred_val, ref_val):
                match_count += 1

    return match_count


_STEREO_ERROR_MARGIN = 0.5

def score_stereoisomer_match(pred_row, gt_row, error_margin: float = _STEREO_ERROR_MARGIN,
                               tiebreak_fields: list = None) -> int:
    """Stereoisomer tiebreak scoring (stage 3)"""
    score = 0
    for field, pts in [('DC50', 2), ('Dmax', 2)]:
        pred_val = pred_row.get(field) if hasattr(pred_row, 'get') else None
        gt_val   = gt_row.get(field)   if hasattr(gt_row, 'get')   else None

        if is_missing_value(pred_val) or is_missing_value(gt_val):
            continue

        _, pred_num = extract_value_with_operator(str(pred_val))
        _, gt_num   = extract_value_with_operator(str(gt_val))

        if pred_num is not None and gt_num is not None:
            if abs(pred_num - gt_num) <= error_margin:
                score += pts
        else:
            pred_norm = normalize_unicode_chars(str(pred_val)).lower()
            gt_norm   = normalize_unicode_chars(str(gt_val)).lower()
            if pred_norm == gt_norm:
                score += pts

    if tiebreak_fields:
        for field in tiebreak_fields:
            pred_val = pred_row.get(field) if hasattr(pred_row, 'get') else None
            gt_val   = gt_row.get(field)   if hasattr(gt_row, 'get')   else None
            if is_missing_value(pred_val) or is_missing_value(gt_val):
                continue
            if compare_exact_fields_with_normalization(pred_val, gt_val):
                score += 1

    return score


def _classify_fp_type(pred_compound: str, attempted_matches: List[Dict],
                      gt_compounds_in_paper: Set[str]) -> str:
    """Classify FP into three types:"""
    if not attempted_matches:
        pred_lower = pred_compound.lower().strip()
        for gt_comp in gt_compounds_in_paper:
            gt_lower = gt_comp.lower().strip()
            if (pred_lower in gt_lower or gt_lower in pred_lower or
                pred_lower.replace(' ', '') == gt_lower.replace(' ', '')):
                return 'compound_name_error'

        return 'compound_not_in_labeled_data'
    else:
        return 'field_mismatch'


def _generate_fp_summary(fp_type: str, pred_row: pd.Series,
                        attempted_matches: List[Dict],
                        gt_compounds_in_paper: Set[str]) -> Dict:
    """Generate FP summary based on FP type classification"""
    pred_compound = str(pred_row.get('Compound_Name', 'unknown'))

    if fp_type == 'compound_not_in_labeled_data':
        return {
            "fp_type": "compound_not_in_labeled_data",
            "quick_explanation": "Incorrectly extracted DC50/Dmax degradation assay measurements: This compound has no DC50 or Dmax degradation assay measurements reported in the paper."
        }

    elif fp_type == 'compound_name_error':

        def clean_compound_name(name):
            """Remove common prefixes like 'Compound', 'Cpd', 'No"""
            return re.sub(r'^(compound|comp|cpd|no\.?)\s*', '', name.lower().strip())

        pred_lower = pred_compound.lower().strip()
        pred_cleaned = clean_compound_name(pred_lower)
        similar_gt_compound = None

        best_match_score = 0
        best_match_comp = None

        for gt_comp in gt_compounds_in_paper:
            gt_lower = gt_comp.lower().strip()
            gt_cleaned = clean_compound_name(gt_lower)

            current_score = 0

            if pred_lower == gt_lower:
                current_score = 3

            elif pred_cleaned == gt_cleaned:
                current_score = 2

            elif re.search(r'\b' + re.escape(gt_lower) + r'\b', pred_lower):
                current_score = 1

            if current_score > best_match_score:
                best_match_score = current_score
                best_match_comp = gt_comp

                if best_match_score == 3:
                    break

        if best_match_comp:
            similar_gt_compound = best_match_comp

        if similar_gt_compound:
            explanation = f"Incorrectly extracted compound name: Ground truth standardised name is {similar_gt_compound}, extracted name is {pred_compound}"
        else:
            explanation = f"Incorrectly extracted compound name: Extracted name {pred_compound} does not match any compound name in ground truth"

        return {
            "fp_type": "compound_name_error",
            "quick_explanation": explanation,
            "gt_candidate": similar_gt_compound
        }

    else:
        if not attempted_matches:
            return {
                "fp_type": "field_mismatch",
                "quick_explanation": "Incorrectly extracted compound information: Field matching failed (reason unknown)"
            }

        mismatch_fields_counter = {}
        all_gt_values_by_field = {}

        for attempt in attempted_matches:
            for field in attempt.get('mismatch_fields', []):
                mismatch_fields_counter[field] = mismatch_fields_counter.get(field, 0) + 1

                if field not in all_gt_values_by_field:
                    all_gt_values_by_field[field] = []

                details = attempt.get('mismatch_details', {}).get(field, {})
                gt_val = details.get('label', '')
                if gt_val and gt_val not in all_gt_values_by_field[field]:
                    all_gt_values_by_field[field].append(gt_val)

        primary_field = max(mismatch_fields_counter.items(), key=lambda x: x[1])[0]
        pred_value = str(pred_row.get(primary_field, ''))
        gt_values = all_gt_values_by_field.get(primary_field, [])

        if len(gt_values) == 1:
            gt_values_str = gt_values[0]
        elif len(gt_values) <= 4:
            gt_values_str = f"[{', '.join(gt_values)}]"
        else:
            gt_values_str = f"[{', '.join(gt_values[:3])}, ... ({len(gt_values)} total)]"

        explanation = f"Incorrectly extracted compound information: Label {primary_field} is {gt_values_str}, extracted {primary_field} is {pred_value}"

        return {
            "fp_type": "field_mismatch",
            "quick_explanation": explanation,
            "mismatched_field": primary_field,
            "pred_value": pred_value,
            "gt_values": gt_values,
            "total_gt_candidates": len(attempted_matches)
        }


def evaluate_record_based_accuracy(pred_csv_path: str, labeled_csv_path: str,
                                    config: dict = None,
                                    llm_api_key: str = None,
                                    llm_api_url: str = None,
                                    cache_file_path: str = None,
                                    lenient_mode: bool = False,
                                    output_dir: str = None,
                                    ternary_complex_level: bool = False) -> dict:
    """Record-based evaluation"""
    print("-" * 60)
    print("Starting record-based evaluation.")

    try:
        if cache_file_path is None:
            if config and "evaluation" in config and "cache_file_path" in config["evaluation"]:
                config_cache_path = config["evaluation"]["cache_file_path"]
                if config_cache_path is not None and config_cache_path.strip():
                    cache_file_path = config_cache_path
                    print(f"Using cache file from config: {cache_file_path}")
                    cache_dir = Path(cache_file_path).parent
                    cache_dir.mkdir(parents=True, exist_ok=True)

            if cache_file_path is None:
                pred_csv_dir = Path(pred_csv_path).parent
                pred_csv_dir.mkdir(parents=True, exist_ok=True)
                cache_file_path = str(pred_csv_dir / "semantic_match_cache.json")
                print(f"Using default per-run cache file: {cache_file_path}")

        pred_df = pd.read_csv(pred_csv_path)
        labeled_df = pd.read_csv(labeled_csv_path)
        labeled_df.columns = labeled_df.columns.str.strip()
        pred_df.columns = pred_df.columns.str.strip()

        filter_dois = None
        if config and "data" in config and "paper_filter_file" in config["data"]:
            paper_filter_file = config["data"]["paper_filter_file"]
            if paper_filter_file and Path(paper_filter_file).exists():
                print(f"\nFound paper_filter_file: {paper_filter_file}")
                try:
                    with open(paper_filter_file, 'r', encoding='utf-8') as f:
                        filter_data = json.load(f)

                    filter_dois = set()
                    for item in filter_data:
                        if 'doi' in item:
                            filter_dois.add(item['doi'])

                    if filter_dois:
                        print(f"Loaded {len(filter_dois)} DOIs from paper_filter_file")
                        print(f"Will filter ground truth to only include these DOIs")

                        gt_records_before = len(labeled_df)
                        gt_dois_before = set(labeled_df['DOI'].unique())

                        labeled_df = labeled_df[labeled_df['DOI'].isin(filter_dois)]

                        gt_records_after = len(labeled_df)
                        gt_dois_after = set(labeled_df['DOI'].unique())

                        print(f"  Ground truth DOIs (before): {len(gt_dois_before)}")
                        print(f"  Ground truth DOIs (after): {len(gt_dois_after)}")
                        print(f"  Ground truth records (before): {gt_records_before}")
                        print(f"  Ground truth records (after): {gt_records_after}")
                        print(f"  Filtered out: {gt_records_before - gt_records_after} records from {len(gt_dois_before) - len(gt_dois_after)} DOIs")
                except Exception as e:
                    print(f"Warning: Could not load paper_filter_file: {e}")
                    print("Proceeding with full ground truth")

        gt_dois = set(labeled_df['DOI'].unique())
        pred_dois_before = set(pred_df['DOI'].unique())
        records_before = len(pred_df)

        pred_df = pred_df[pred_df['DOI'].isin(gt_dois)]

        records_after = len(pred_df)
        pred_dois_after = set(pred_df['DOI'].unique())

        print(f"\n{'='*60}")
        print(f"DOI Filtering (Predictions):")
        print(f"  Ground truth DOIs: {len(gt_dois)}")
        print(f"  Prediction DOIs (before): {len(pred_dois_before)}")
        print(f"  Prediction DOIs (after): {len(pred_dois_after)}")
        print(f"  Prediction records (before): {records_before}")
        print(f"  Prediction records (after): {records_after}")
        print(f"  Filtered out: {records_before - records_after} records from {len(pred_dois_before) - len(pred_dois_after)} DOIs")
        print(f"{'='*60}")

        if ternary_complex_level:
            record_fields = ['Compound_Name', 'Degradation_Target', 'Recruiter']
            semantic_record_fields = ['Degradation_Target', 'Recruiter']
            assay_fields = ['Assay', 'Cell_Line', 'DC50', 'DC50_units', 'DC50_h', 'Dmax', 'Dmax_h', 'Dmax_conc']
            tiebreak_fields_for_score = ['Cell_Line', 'Assay']
        else:
            record_fields = ['Compound_Name', 'Degradation_Target', 'Recruiter', 'Assay', 'Cell_Line']
            semantic_record_fields = ['Degradation_Target', 'Recruiter', 'Assay', 'Cell_Line']
            assay_fields = ['DC50', 'DC50_units', 'DC50_h', 'Dmax', 'Dmax_h', 'Dmax_conc']
            tiebreak_fields_for_score = None

        print("\n" + "=" * 60)
        print("Part 1: Record-level Evaluation")
        print("=" * 60)

        part1_results = {
            'total_tp': 0,
            'total_fp': 0,
            'total_fn': 0,
            'per_doi_metrics': {},
            'matched_records': []
        }

        detailed_results_by_doi = {}

        gt_compounds_by_doi = {}
        for doi in labeled_df['DOI'].unique():
            doi_labeled_df = labeled_df[labeled_df['DOI'] == doi]
            gt_compounds_by_doi[doi] = set(doi_labeled_df['Compound_Name'].dropna().unique())

        need_review_records = []

        for doi in labeled_df['DOI'].unique():
            doi_pred_df = pred_df[pred_df['DOI'] == doi]
            doi_labeled_df = labeled_df[labeled_df['DOI'] == doi]

            doi_tp = 0
            doi_fp = 0
            doi_fn = 0

            successfully_matched_label_indices = set()

            detailed_results_by_doi[str(doi)] = {
                'part1_record_level': {
                    'tp_records': [],
                    'fp_records': [],
                    'fn_records': [],
                    'need_review_records': []
                },
                'part2_field_level': {
                    'matched_records': []
                }
            }

            pred_candidates_map = {}
            pred_compounds_map = {}

            for pred_idx, pred_row in doi_pred_df.iterrows():
                pred_compound = pred_row['Compound_Name'] if 'Compound_Name' in pred_row else None
                pred_conn_key_raw = pred_row['Connectivity_Key'] if 'Connectivity_Key' in pred_row else None
                has_compound_name = not pd.isna(pred_compound)
                has_conn_key = pd.notna(pred_conn_key_raw) and str(pred_conn_key_raw).strip().lower() not in ('', 'nan')

                if not has_compound_name and not has_conn_key:
                    print(f"Warning: Skipping pred row {pred_idx} with missing Compound_Name and Connectivity_Key")
                    continue

                if not has_compound_name:
                    pred_compound = f"[conn_key:{str(pred_conn_key_raw).strip()}]"

                matched_by_conn_key = False

                if has_compound_name:
                    pred_compound_norm = normalize_dashes(str(pred_compound).lower())
                    matching_labeled = doi_labeled_df[
                        doi_labeled_df['Compound_Name'].apply(lambda x: normalize_dashes(str(x).lower())) == pred_compound_norm
                    ]
                else:
                    matching_labeled = doi_labeled_df.iloc[0:0]

                if matching_labeled.empty and has_conn_key and 'Connectivity_Key' in doi_labeled_df.columns:
                    pred_conn_key = normalize_dashes(str(pred_conn_key_raw).strip().lower())
                    matching_labeled = doi_labeled_df[
                        doi_labeled_df['Connectivity_Key'].apply(lambda x: normalize_dashes(str(x).lower())) == pred_conn_key
                    ]
                    if not matching_labeled.empty:
                        matched_by_conn_key = True

                if matching_labeled.empty:
                    doi_fp += 1

                    fp_record = {
                        'pred_compound': str(pred_compound),
                        'pred_target': str(pred_row.get('Degradation_Target', '')),
                        'pred_recruiter': str(pred_row.get('Recruiter', '')),
                        'pred_assay': str(pred_row.get('Assay', '')),
                        'pred_cell_line': str(pred_row.get('Cell_Line', '')),
                        'fp_reason': 'Compound not in ground truth',
                        'attempted_matches': []
                    }

                    gt_compounds = gt_compounds_by_doi.get(doi, set())
                    fp_type = _classify_fp_type(str(pred_compound), [], gt_compounds)
                    fp_summary = _generate_fp_summary(fp_type, pred_row, [], gt_compounds)
                    fp_record['fp_summary'] = fp_summary

                    detailed_results_by_doi[str(doi)]['part1_record_level']['fp_records'].append(fp_record)

                    continue

                candidates = []
                attempted_matches = []

                for label_idx, label_row in matching_labeled.iterrows():
                    if label_idx in successfully_matched_label_indices:
                        continue

                    mismatch_info = {
                        'label_record_id': label_row.get('Record_ID', None),
                        'label_target': str(label_row.get('Degradation_Target', '')),
                        'label_recruiter': str(label_row.get('Recruiter', '')),
                        'label_assay': str(label_row.get('Assay', '')),
                        'label_cell_line': str(label_row.get('Cell_Line', '')),
                        'mismatch_fields': [],
                        'mismatch_details': {}
                    }

                    all_fields_match = True

                    for field in ['Degradation_Target', 'Recruiter']:
                        if field not in pred_row or field not in label_row:
                            all_fields_match = False
                            mismatch_info['mismatch_fields'].append(field)
                            mismatch_info['mismatch_details'][field] = {
                                'pred': str(pred_row.get(field, '')),
                                'label': str(label_row.get(field, '')),
                                'reason': 'Field missing in data'
                            }
                            break

                        pred_val = pred_row[field]
                        label_val = label_row[field]

                        if (field == 'Recruiter'
                                and not is_missing_value(label_val)
                                and not is_missing_value(pred_val)
                                and str(pred_val).strip().lower() == 'unknown'):
                            all_fields_match = False
                            mismatch_info['mismatch_fields'].append(field)
                            mismatch_info['mismatch_details'][field] = {
                                'pred': str(pred_val),
                                'label': str(label_val),
                                'reason': 'Pred Recruiter is Unknown'
                            }
                            break

                        is_match = check_semantic_match(
                            pred_val, label_val, feature=field,
                            llm_api_key=llm_api_key, llm_api_url=llm_api_url,
                            config=config, cache_file_path=cache_file_path,
                            doi=doi
                        )
                        if not is_match:
                            all_fields_match = False
                            mismatch_info['mismatch_fields'].append(field)
                            mismatch_info['mismatch_details'][field] = {
                                'pred': str(pred_val),
                                'label': str(label_val),
                                'reason': 'Semantic match failed'
                            }
                            break

                    if not all_fields_match:
                        attempted_matches.append(mismatch_info)
                        continue

                    assay_cell_match = True
                    ref_has_empty_assay_or_cell = False

                    if not ternary_complex_level:
                        for field in ['Assay', 'Cell_Line']:
                            if field not in pred_row or field not in label_row:
                                continue

                            pred_val = pred_row[field]
                            label_val = label_row[field]

                            if is_missing_value(label_val):
                                ref_has_empty_assay_or_cell = True
                                continue

                            is_match = check_semantic_match(
                                pred_val, label_val, feature=field,
                                llm_api_key=llm_api_key, llm_api_url=llm_api_url,
                                config=config, cache_file_path=cache_file_path,
                                doi=doi
                            )
                            if not is_match:
                                assay_cell_match = False
                                mismatch_info['mismatch_fields'].append(field)
                                mismatch_info['mismatch_details'][field] = {
                                    'pred': str(pred_val),
                                    'label': str(label_val),
                                    'reason': 'Semantic match failed'
                                }
                                break

                    if assay_cell_match:
                        candidates.append((label_idx, label_row, ref_has_empty_assay_or_cell))
                    else:
                        attempted_matches.append(mismatch_info)

                pred_candidates_map[pred_idx] = {
                    'pred_row': pred_row,
                    'candidates': candidates,
                    'attempted_matches': attempted_matches,
                    'matched_by_conn_key': matched_by_conn_key,
                }
                pred_compounds_map[pred_idx] = pred_compound

            ref_to_competing_preds = {}
            for pred_idx, pred_info in pred_candidates_map.items():
                for label_idx, label_row, ref_has_empty in pred_info['candidates']:
                    if label_idx in successfully_matched_label_indices:
                        continue
                    if label_idx not in ref_to_competing_preds:
                        ref_to_competing_preds[label_idx] = []
                    ref_to_competing_preds[label_idx].append(pred_idx)

            successfully_matched_pred_indices = set()

            for label_idx, competing_pred_indices in ref_to_competing_preds.items():
                if label_idx in successfully_matched_label_indices:
                    continue

                available_competing_preds = [
                    p for p in competing_pred_indices
                    if p not in successfully_matched_pred_indices
                ]

                if len(available_competing_preds) < 2:
                    continue

                label_row = doi_labeled_df.loc[label_idx]
                best_pred_idx = None
                best_score = -1

                for pred_idx in available_competing_preds:
                    pred_row = pred_candidates_map[pred_idx]['pred_row']
                    score = calculate_assay_match_score(
                        pred_row, label_row,
                        tiebreak_fields=tiebreak_fields_for_score,
                    )
                    if score > best_score:
                        best_score = score
                        best_pred_idx = pred_idx

                if best_pred_idx is not None:
                    successfully_matched_label_indices.add(label_idx)
                    successfully_matched_pred_indices.add(best_pred_idx)
                    pred_row = pred_candidates_map[best_pred_idx]['pred_row']
                    pred_compound = pred_compounds_map[best_pred_idx]
                    matched_record_id = label_row['Record_ID'] if 'Record_ID' in label_row else None

                    part1_results['matched_records'].append({
                        'pred_index': best_pred_idx,
                        'label_index': label_idx,
                        'record_id': matched_record_id,
                        'doi': doi,
                        'compound_name': pred_compound
                    })
                    doi_tp += 1

                    tp_record = {
                        'pred_compound': str(pred_compound),
                        'pred_target': str(pred_row.get('Degradation_Target', '')),
                        'pred_recruiter': str(pred_row.get('Recruiter', '')),
                        'pred_assay': str(pred_row.get('Assay', '')),
                        'pred_cell_line': str(pred_row.get('Cell_Line', '')),
                        'label_record_id': matched_record_id,
                        'label_compound': str(label_row.get('Compound_Name', '')),
                        'label_target': str(label_row.get('Degradation_Target', '')),
                        'label_recruiter': str(label_row.get('Recruiter', '')),
                        'label_assay': str(label_row.get('Assay', '')),
                        'label_cell_line': str(label_row.get('Cell_Line', '')),
                        'match_reason': f'Selected from {len(available_competing_preds)} competing predictions (best assay match score: {best_score})'
                    }
                    detailed_results_by_doi[str(doi)]['part1_record_level']['tp_records'].append(tp_record)

            for pred_idx, pred_info in pred_candidates_map.items():
                if pred_idx in successfully_matched_pred_indices:
                    continue
                pred_row = pred_info['pred_row']
                candidates = pred_info['candidates']
                attempted_matches = pred_info['attempted_matches']
                pred_compound = pred_compounds_map[pred_idx]

                available_candidates = [
                    (label_idx, label_row, ref_has_empty)
                    for label_idx, label_row, ref_has_empty in candidates
                    if label_idx not in successfully_matched_label_indices
                ]

                found_match = False
                matched_label_idx = None

                if len(available_candidates) == 0:
                    doi_fp += 1

                    fp_record = {
                        'pred_compound': str(pred_compound),
                        'pred_target': str(pred_row.get('Degradation_Target', '')),
                        'pred_recruiter': str(pred_row.get('Recruiter', '')),
                        'pred_assay': str(pred_row.get('Assay', '')),
                        'pred_cell_line': str(pred_row.get('Cell_Line', '')),
                        'fp_reason': 'No matching record found (all candidate records failed field matching)',
                        'attempted_matches': attempted_matches
                    }

                    gt_compounds = gt_compounds_by_doi.get(doi, set())
                    fp_type = _classify_fp_type(str(pred_compound), attempted_matches, gt_compounds)
                    fp_summary = _generate_fp_summary(fp_type, pred_row, attempted_matches, gt_compounds)
                    fp_record['fp_summary'] = fp_summary

                    detailed_results_by_doi[str(doi)]['part1_record_level']['fp_records'].append(fp_record)

                elif len(available_candidates) == 1:
                    label_idx, label_row, ref_has_empty = available_candidates[0]

                    competing_preds = ref_to_competing_preds.get(label_idx, [])
                    has_competition = len(competing_preds) > 1

                    pred_has_all_fields = all(
                        not is_missing_value(pred_row.get(f))
                        for f in ['Compound_Name', 'Degradation_Target', 'Recruiter', 'Assay', 'Cell_Line']
                    )

                    if has_competition and ref_has_empty and pred_has_all_fields:
                        assay_score = calculate_assay_match_score(
                            pred_row, label_row,
                            tiebreak_fields=tiebreak_fields_for_score,
                        )
                        if assay_score > 0:
                            found_match = True
                            matched_label_idx = label_idx
                            successfully_matched_label_indices.add(label_idx)
                        else:
                            pass
                    else:
                        found_match = True
                        matched_label_idx = label_idx
                        successfully_matched_label_indices.add(label_idx)

                    if found_match:
                        matched_record_id = label_row['Record_ID'] if 'Record_ID' in label_row else None
                        part1_results['matched_records'].append({
                            'pred_index': pred_idx,
                            'label_index': label_idx,
                            'record_id': matched_record_id,
                            'doi': doi,
                            'compound_name': pred_compound
                        })
                        doi_tp += 1

                        tp_record = {
                            'pred_compound': str(pred_compound),
                            'pred_target': str(pred_row.get('Degradation_Target', '')),
                            'pred_recruiter': str(pred_row.get('Recruiter', '')),
                            'pred_assay': str(pred_row.get('Assay', '')),
                            'pred_cell_line': str(pred_row.get('Cell_Line', '')),
                            'label_record_id': matched_record_id,
                            'label_compound': str(label_row.get('Compound_Name', '')),
                            'label_target': str(label_row.get('Degradation_Target', '')),
                            'label_recruiter': str(label_row.get('Recruiter', '')),
                            'label_assay': str(label_row.get('Assay', '')),
                            'label_cell_line': str(label_row.get('Cell_Line', '')),
                            'match_reason': f'All 5 record fields matched (1 candidate, {"with competition resolved by assay values" if has_competition else "no competition"})'
                        }
                        detailed_results_by_doi[str(doi)]['part1_record_level']['tp_records'].append(tp_record)

                else:
                    pred_matched_by_conn_key = pred_info.get('matched_by_conn_key', False)

                    if pred_matched_by_conn_key:
                        best_candidate = None
                        max_score = -1

                        for label_idx, label_row, ref_has_empty in available_candidates:
                            current_score = score_stereoisomer_match(
                                pred_row, label_row,
                                tiebreak_fields=tiebreak_fields_for_score,
                            )
                            if current_score > max_score:
                                max_score = current_score
                                best_candidate = (label_idx, label_row)

                        if max_score > 0:
                            label_idx, label_row = best_candidate
                            found_match = True
                            matched_label_idx = label_idx
                            successfully_matched_label_indices.add(label_idx)
                            matched_record_id = label_row['Record_ID'] if 'Record_ID' in label_row else None
                            part1_results['matched_records'].append({
                                'pred_index': pred_idx,
                                'label_index': label_idx,
                                'record_id': matched_record_id,
                                'doi': doi,
                                'compound_name': pred_compound
                            })
                            doi_tp += 1
                            tp_record = {
                                'pred_compound': str(pred_compound),
                                'pred_target': str(pred_row.get('Degradation_Target', '')),
                                'pred_recruiter': str(pred_row.get('Recruiter', '')),
                                'pred_assay': str(pred_row.get('Assay', '')),
                                'pred_cell_line': str(pred_row.get('Cell_Line', '')),
                                'label_record_id': matched_record_id,
                                'label_compound': str(label_row.get('Compound_Name', '')),
                                'label_target': str(label_row.get('Degradation_Target', '')),
                                'label_recruiter': str(label_row.get('Recruiter', '')),
                                'label_assay': str(label_row.get('Assay', '')),
                                'label_cell_line': str(label_row.get('Cell_Line', '')),
                                'match_reason': f'Connectivity_Key stereoisomer tiebreak (score={max_score}, {len(available_candidates)} candidates)'
                            }
                            detailed_results_by_doi[str(doi)]['part1_record_level']['tp_records'].append(tp_record)

                        else:
                            for label_idx, label_row, _ in available_candidates:
                                successfully_matched_label_indices.add(label_idx)
                            found_match = True

                            nr_record = {
                                'pred_compound': str(pred_compound),
                                'pred_dc50': str(pred_row.get('DC50', '')),
                                'pred_dmax': str(pred_row.get('Dmax', '')),
                                'gt_candidates': [
                                    {
                                        'compound_name': str(lr.get('Compound_Name', '')),
                                        'record_id': str(lr.get('Record_ID', '')),
                                        'dc50': str(lr.get('DC50', '')),
                                        'dmax': str(lr.get('Dmax', ''))
                                    }
                                    for _, lr, _ in available_candidates
                                ],
                                'reason': 'Connectivity_Key 1-to-many: stereoisomers indistinguishable by activity data'
                            }
                            detailed_results_by_doi[str(doi)]['part1_record_level']['need_review_records'].append(nr_record)
                            need_review_records.append({'doi': doi, **nr_record})

                    else:
                        best_candidate = None
                        best_score = -1

                        for label_idx, label_row, ref_has_empty in available_candidates:
                            score = calculate_assay_match_score(
                                pred_row, label_row,
                                tiebreak_fields=tiebreak_fields_for_score,
                            )
                            if score > best_score:
                                best_score = score
                                best_candidate = (label_idx, label_row)

                        if best_candidate is None:
                            best_candidate = (available_candidates[0][0], available_candidates[0][1])

                        label_idx, label_row = best_candidate
                        found_match = True
                        matched_label_idx = label_idx
                        successfully_matched_label_indices.add(label_idx)

                        matched_record_id = label_row['Record_ID'] if 'Record_ID' in label_row else None
                        part1_results['matched_records'].append({
                            'pred_index': pred_idx,
                            'label_index': label_idx,
                            'record_id': matched_record_id,
                            'doi': doi,
                            'compound_name': pred_compound
                        })
                        doi_tp += 1

                        tp_record = {
                            'pred_compound': str(pred_compound),
                            'pred_target': str(pred_row.get('Degradation_Target', '')),
                            'pred_recruiter': str(pred_row.get('Recruiter', '')),
                            'pred_assay': str(pred_row.get('Assay', '')),
                            'pred_cell_line': str(pred_row.get('Cell_Line', '')),
                            'label_record_id': matched_record_id,
                            'label_compound': str(label_row.get('Compound_Name', '')),
                            'label_target': str(label_row.get('Degradation_Target', '')),
                            'label_recruiter': str(label_row.get('Recruiter', '')),
                            'label_assay': str(label_row.get('Assay', '')),
                            'label_cell_line': str(label_row.get('Cell_Line', '')),
                            'match_reason': f'All 5 record fields matched (selected from {len(available_candidates)} candidates using assay value matching)'
                        }
                        detailed_results_by_doi[str(doi)]['part1_record_level']['tp_records'].append(tp_record)

                if not found_match and len(available_candidates) > 0:
                    doi_fp += 1
                    fp_record = {
                        'pred_compound': str(pred_compound),
                        'pred_target': str(pred_row.get('Degradation_Target', '')),
                        'pred_recruiter': str(pred_row.get('Recruiter', '')),
                        'pred_assay': str(pred_row.get('Assay', '')),
                        'pred_cell_line': str(pred_row.get('Cell_Line', '')),
                        'fp_reason': 'Candidate found but assay values did not match (failed competition)',
                        'attempted_matches': attempted_matches
                    }
                    detailed_results_by_doi[str(doi)]['part1_record_level']['fp_records'].append(fp_record)

            if doi_pred_df.empty:
                predicted_compounds = set()
            else:
                predicted_compounds = set(doi_pred_df['Compound_Name'].apply(lambda x: normalize_dashes(str(x).lower())))

            for label_idx, label_row in doi_labeled_df.iterrows():
                if label_idx in successfully_matched_label_indices:
                    continue

                label_compound = label_row['Compound_Name']
                compound_in_pred = normalize_dashes(str(label_compound).lower()) in predicted_compounds

                if not compound_in_pred:
                    doi_fn += 1

                    fn_record = {
                        'label_record_id': label_row.get('Record_ID', None),
                        'label_compound': str(label_compound),
                        'label_target': str(label_row.get('Degradation_Target', '')),
                        'label_recruiter': str(label_row.get('Recruiter', '')),
                        'label_assay': str(label_row.get('Assay', '')),
                        'label_cell_line': str(label_row.get('Cell_Line', '')),
                        'fn_reason': 'Compound not in predictions'
                    }
                    detailed_results_by_doi[str(doi)]['part1_record_level']['fn_records'].append(fn_record)
                else:
                    doi_fn += 1

                    fn_record = {
                        'label_record_id': label_row.get('Record_ID', None),
                        'label_compound': str(label_compound),
                        'label_target': str(label_row.get('Degradation_Target', '')),
                        'label_recruiter': str(label_row.get('Recruiter', '')),
                        'label_assay': str(label_row.get('Assay', '')),
                        'label_cell_line': str(label_row.get('Cell_Line', '')),
                        'fn_reason': 'Compound exists in predictions but no record matched all 5 fields'
                    }
                    detailed_results_by_doi[str(doi)]['part1_record_level']['fn_records'].append(fn_record)

            doi_precision = doi_tp / (doi_tp + doi_fp) if (doi_tp + doi_fp) > 0 else 0
            doi_recall = doi_tp / (doi_tp + doi_fn) if (doi_tp + doi_fn) > 0 else 0
            doi_f1 = 2 * doi_precision * doi_recall / (doi_precision + doi_recall) if (doi_precision + doi_recall) > 0 else 0

            part1_results['per_doi_metrics'][str(doi)] = {
                'tp': int(doi_tp),
                'fp': int(doi_fp),
                'fn': int(doi_fn),
                'precision': float(doi_precision),
                'recall': float(doi_recall),
                'f1_score': float(doi_f1)
            }

            part1_results['total_tp'] += doi_tp
            part1_results['total_fp'] += doi_fp
            part1_results['total_fn'] += doi_fn

        total_tp = part1_results['total_tp']
        total_fp = part1_results['total_fp']
        total_fn = part1_results['total_fn']

        precision_records_micro = total_tp / (total_tp + total_fp) if (total_tp + total_fp) > 0 else 0
        recall_records_micro = total_tp / (total_tp + total_fn) if (total_tp + total_fn) > 0 else 0
        f1_records_micro = 2 * precision_records_micro * recall_records_micro / (precision_records_micro + recall_records_micro) if (precision_records_micro + recall_records_micro) > 0 else 0

        per_doi_precisions = []
        per_doi_recalls = []
        per_doi_f1s = []

        for doi, metrics in part1_results['per_doi_metrics'].items():
            per_doi_precisions.append(metrics['precision'])
            per_doi_recalls.append(metrics['recall'])
            per_doi_f1s.append(metrics['f1_score'])

        precision_records = sum(per_doi_precisions) / len(per_doi_precisions) if per_doi_precisions else 0
        recall_records = sum(per_doi_recalls) / len(per_doi_recalls) if per_doi_recalls else 0
        f1_records = sum(per_doi_f1s) / len(per_doi_f1s) if per_doi_f1s else 0

        print(f"\nPart 1 Results:")
        print(f"  Total TP (correctly identified records): {total_tp}")
        print(f"  Total FP (incorrectly identified records): {total_fp}")
        print(f"  Total FN (missed records): {total_fn}")
        print(f"  Precision (Records, Macro-avg): {precision_records:.4f}")
        print(f"  Recall (Records, Macro-avg): {recall_records:.4f}")
        print(f"  F1-score (Records, Macro-avg): {f1_records:.4f}")
        print(f"  Precision (Records, Micro-avg): {precision_records_micro:.4f}")
        print(f"  Recall (Records, Micro-avg): {recall_records_micro:.4f}")
        print(f"  F1-score (Records, Micro-avg): {f1_records_micro:.4f}")

        print("\n" + "=" * 60)
        print("Part 2: Entry-level Evaluation (for TP records only)")
        print("=" * 60)

        part2_results = {
            'total_tp': 0,
            'total_fp': 0,
            'total_fn': 0,
            'total_tn': 0,
            'per_field_metrics': {},
            'per_doi_field_metrics': {},
            'entry_details': []
        }

        for field in assay_fields:
            part2_results['per_field_metrics'][field] = {
                'tp': 0,
                'fp': 0,
                'fn': 0,
                'tn': 0
            }

        for matched_record in part1_results['matched_records']:
            pred_idx = matched_record['pred_index']
            label_idx = matched_record['label_index']
            record_id = matched_record['record_id']
            doi = matched_record['doi']

            if str(doi) not in part2_results['per_doi_field_metrics']:
                part2_results['per_doi_field_metrics'][str(doi)] = {
                    'total_tp': 0,
                    'total_fp': 0,
                    'total_fn': 0,
                    'total_tn': 0,
                    'per_field': {}
                }
                for field in assay_fields:
                    part2_results['per_doi_field_metrics'][str(doi)]['per_field'][field] = {
                        'tp': 0,
                        'fp': 0,
                        'fn': 0,
                        'tn': 0
                    }

            pred_row = pred_df.loc[pred_idx]
            label_row = labeled_df.loc[label_idx]

            record_fields_detail = {
                'record_id': record_id,
                'compound': str(matched_record['compound_name']),
                'fields': {}
            }

            for field in assay_fields:
                if field not in pred_row or field not in label_row:
                    continue

                pred_val = pred_row[field]
                label_val = label_row[field]

                pred_is_missing = is_missing_value(pred_val)
                label_is_missing = is_missing_value(label_val)

                field_detail = {
                    'pred_value': str(pred_val) if not pred_is_missing else '',
                    'label_value': str(label_val) if not label_is_missing else ''
                }

                if pred_is_missing and label_is_missing:
                    part2_results['total_tn'] += 1
                    part2_results['per_field_metrics'][field]['tn'] += 1
                    part2_results['per_doi_field_metrics'][str(doi)]['total_tn'] += 1
                    part2_results['per_doi_field_metrics'][str(doi)]['per_field'][field]['tn'] += 1

                    field_detail['status'] = 'TN'
                    field_detail['match'] = None
                    field_detail['reason'] = 'Both empty'
                elif pred_is_missing and not label_is_missing:
                    part2_results['total_fn'] += 1
                    part2_results['per_field_metrics'][field]['fn'] += 1
                    part2_results['per_doi_field_metrics'][str(doi)]['total_fn'] += 1
                    part2_results['per_doi_field_metrics'][str(doi)]['per_field'][field]['fn'] += 1

                    field_detail['status'] = 'FN'
                    field_detail['match'] = False
                    field_detail['reason'] = 'Missing extraction'
                elif not pred_is_missing and label_is_missing:
                    part2_results['total_fp'] += 1
                    part2_results['per_field_metrics'][field]['fp'] += 1
                    part2_results['per_doi_field_metrics'][str(doi)]['total_fp'] += 1
                    part2_results['per_doi_field_metrics'][str(doi)]['per_field'][field]['fp'] += 1

                    field_detail['status'] = 'FP'
                    field_detail['match'] = False
                    field_detail['reason'] = 'Extra prediction (label empty)'
                else:
                    if field in ('Assay', 'Cell_Line'):
                        values_match = check_semantic_match(
                            pred_val, label_val, feature=field,
                            llm_api_key=llm_api_key, llm_api_url=llm_api_url,
                            config=config, cache_file_path=cache_file_path,
                            doi=doi
                        )
                    elif ternary_complex_level and field == 'DC50_units':
                        values_match = compare_concentration_units(pred_val, label_val)
                    else:
                        values_match = compare_assay_values(pred_val, label_val)
                    if values_match:
                        part2_results['total_tp'] += 1
                        part2_results['per_field_metrics'][field]['tp'] += 1
                        part2_results['per_doi_field_metrics'][str(doi)]['total_tp'] += 1
                        part2_results['per_doi_field_metrics'][str(doi)]['per_field'][field]['tp'] += 1

                        field_detail['status'] = 'TP'
                        field_detail['match'] = True
                        field_detail['reason'] = 'Values matched'
                    else:
                        part2_results['total_fp'] += 1
                        part2_results['per_field_metrics'][field]['fp'] += 1
                        part2_results['per_doi_field_metrics'][str(doi)]['total_fp'] += 1
                        part2_results['per_doi_field_metrics'][str(doi)]['per_field'][field]['fp'] += 1

                        field_detail['status'] = 'FP'
                        field_detail['match'] = False
                        field_detail['reason'] = 'Predicted wrong value'

                record_fields_detail['fields'][field] = field_detail

                if field in ('Assay', 'Cell_Line'):
                    entry_match = check_semantic_match(
                        pred_val, label_val, feature=field,
                        llm_api_key=llm_api_key, llm_api_url=llm_api_url,
                        config=config, cache_file_path=cache_file_path,
                        doi=doi
                    ) if not (pred_is_missing or label_is_missing) else None
                elif ternary_complex_level and field == 'DC50_units':
                    entry_match = compare_concentration_units(pred_val, label_val) if not (pred_is_missing or label_is_missing) else None
                else:
                    entry_match = compare_assay_values(pred_val, label_val) if not (pred_is_missing or label_is_missing) else None
                part2_results['entry_details'].append({
                    'record_id': record_id,
                    'compound_name': matched_record['compound_name'],
                    'field': field,
                    'pred_value': pred_val,
                    'label_value': label_val,
                    'match': entry_match
                })

            detailed_results_by_doi[str(doi)]['part2_field_level']['matched_records'].append(record_fields_detail)

        total_tp_entry = part2_results['total_tp']
        total_fp_entry = part2_results['total_fp']
        total_fn_entry = part2_results['total_fn']
        total_tn_entry = part2_results['total_tn']

        precision_entry_micro = total_tp_entry / (total_tp_entry + total_fp_entry) if (total_tp_entry + total_fp_entry) > 0 else 0
        recall_entry_micro = total_tp_entry / (total_tp_entry + total_fn_entry) if (total_tp_entry + total_fn_entry) > 0 else 0
        f1_entry_micro = 2 * precision_entry_micro * recall_entry_micro / (precision_entry_micro + recall_entry_micro) if (precision_entry_micro + recall_entry_micro) > 0 else 0

        for field in assay_fields:
            field_metrics = part2_results['per_field_metrics'][field]
            field_tp = field_metrics['tp']
            field_fp = field_metrics['fp']
            field_fn = field_metrics['fn']
            field_tn = field_metrics['tn']

            field_precision = field_tp / (field_tp + field_fp) if (field_tp + field_fp) > 0 else 0
            field_recall = field_tp / (field_tp + field_fn) if (field_tp + field_fn) > 0 else 0
            field_f1 = 2 * field_precision * field_recall / (field_precision + field_recall) if (field_precision + field_recall) > 0 else 0

            part2_results['per_field_metrics'][field]['precision'] = float(field_precision)
            part2_results['per_field_metrics'][field]['recall'] = float(field_recall)
            part2_results['per_field_metrics'][field]['f1_score'] = float(field_f1)

        for doi in part2_results['per_doi_field_metrics']:
            doi_metrics = part2_results['per_doi_field_metrics'][doi]

            doi_tp = doi_metrics['total_tp']
            doi_fp = doi_metrics['total_fp']
            doi_fn = doi_metrics['total_fn']
            doi_tn = doi_metrics['total_tn']

            doi_precision = doi_tp / (doi_tp + doi_fp) if (doi_tp + doi_fp) > 0 else 0
            doi_recall = doi_tp / (doi_tp + doi_fn) if (doi_tp + doi_fn) > 0 else 0
            doi_f1 = 2 * doi_precision * doi_recall / (doi_precision + doi_recall) if (doi_precision + doi_recall) > 0 else 0

            doi_metrics['precision'] = float(doi_precision)
            doi_metrics['recall'] = float(doi_recall)
            doi_metrics['f1_score'] = float(doi_f1)

        per_doi_entry_precisions = []
        per_doi_entry_recalls = []
        per_doi_entry_f1s = []

        for doi, metrics in part2_results['per_doi_field_metrics'].items():
            per_doi_entry_precisions.append(metrics['precision'])
            per_doi_entry_recalls.append(metrics['recall'])
            per_doi_entry_f1s.append(metrics['f1_score'])

        precision_entry = sum(per_doi_entry_precisions) / len(per_doi_entry_precisions) if per_doi_entry_precisions else 0
        recall_entry = sum(per_doi_entry_recalls) / len(per_doi_entry_recalls) if per_doi_entry_recalls else 0
        f1_entry = sum(per_doi_entry_f1s) / len(per_doi_entry_f1s) if per_doi_entry_f1s else 0

        print(f"\nPart 2 Results:")
        print(f"  Total TP (correct entries): {total_tp_entry}")
        print(f"  Total FP (incorrect entries): {total_fp_entry}")
        print(f"  Total FN (missed entries): {total_fn_entry}")
        print(f"  Total TN (both empty): {total_tn_entry}")
        print(f"  Precision (Entry, Macro-avg): {precision_entry:.4f}")
        print(f"  Recall (Entry, Macro-avg): {recall_entry:.4f}")
        print(f"  F1-score (Entry, Macro-avg): {f1_entry:.4f}")
        print(f"  Precision (Entry, Micro-avg): {precision_entry_micro:.4f}")
        print(f"  Recall (Entry, Micro-avg): {recall_entry_micro:.4f}")
        print(f"  F1-score (Entry, Micro-avg): {f1_entry_micro:.4f}")

        final_results = {
            'part1_record_level': {
                'total_tp': int(total_tp),
                'total_fp': int(total_fp),
                'total_fn': int(total_fn),
                'precision_records': float(precision_records),
                'recall_records': float(recall_records),
                'f1_score_records': float(f1_records),
                'precision_records_micro': float(precision_records_micro),
                'recall_records_micro': float(recall_records_micro),
                'f1_score_records_micro': float(f1_records_micro),
                'per_doi_metrics': part1_results['per_doi_metrics']
            },
            'part2_entry_level': {
                'total_tp': int(total_tp_entry),
                'total_fp': int(total_fp_entry),
                'total_fn': int(total_fn_entry),
                'total_tn': int(total_tn_entry),
                'precision_entry': float(precision_entry),
                'recall_entry': float(recall_entry),
                'f1_score_entry': float(f1_entry),
                'precision_entry_micro': float(precision_entry_micro),
                'recall_entry_micro': float(recall_entry_micro),
                'f1_score_entry_micro': float(f1_entry_micro),
                'per_field_metrics': part2_results['per_field_metrics'],
                'per_doi_field_metrics': part2_results['per_doi_field_metrics']
            }
        }

        save_dir = Path(output_dir) if output_dir else Path(pred_csv_path).parent
        save_dir.mkdir(parents=True, exist_ok=True)

        results_path = save_dir / "record_based_evaluation_results.json"
        with open(results_path, "w", encoding="utf-8") as f:
            json.dump(final_results, f, ensure_ascii=False, indent=2)
        print(f"\nRecord-based evaluation results saved to: {results_path}")

        detailed_results_path = save_dir / "evaluation_detailed_by_doi.json"
        with open(detailed_results_path, "w", encoding="utf-8") as f:
            json.dump(detailed_results_by_doi, f, ensure_ascii=False, indent=2)
        print(f"Detailed evaluation results (with per-record labels) saved to: {detailed_results_path}")

        need_review_path = save_dir / "need_review.json"
        with open(need_review_path, "w", encoding="utf-8") as f:
            json.dump(need_review_records, f, ensure_ascii=False, indent=2)
        print(f"Need-review records saved to {need_review_path}")

        print("\n" + "=" * 60)
        print("Record-based evaluation completed successfully!")
        print("=" * 60)

        return final_results

    except Exception as e:
        print(f"Error in record-based evaluation: {e}")
        import traceback
        traceback.print_exc()
        return {
            'error': str(e),
            'part1_record_level': {
                'precision_records': 0,
                'recall_records': 0,
                'f1_score_records': 0
            },
            'part2_entry_level': {
                'precision_entry': 0,
                'recall_entry': 0,
                'f1_score_entry': 0
            }
        }


def run_evaluation(config_file: dict, pred_csv_path: str = None, labeled_csv_path: str = None, lenient_mode: bool = False, output_dir: str = None, ternary_complex_level: bool = False):

    print("-" * 60)
    print("Starting the evaluation process...")

    if not pred_csv_path:
        json_file_path = Path(config_file["results"]["results_dir"])
        all_subdirs = [
            d for d in Path(json_file_path).iterdir()
            if d.is_dir() and len(d.name) == 11
            and d.name[:6].isdigit() and d.name[7:11].isdigit()
        ]
        if not all_subdirs:
            print("No result directory found. Please run the model extraction first.")
            return None

        latest_json_dir = sorted(all_subdirs, key=lambda d: d.name, reverse=True)[0]
        processed_csv_files = list(latest_json_dir.glob("*_processed.csv"))
        if not processed_csv_files:
            print(f"No processed CSV files found in {latest_json_dir}.")
            return None

        pred_csv_path = str(processed_csv_files[0])
        print(f"Using the latest processed CSV file: {pred_csv_path}")

    if not labeled_csv_path:
        if "evaluation" in config_file and "labeled_data_path" in config_file["evaluation"]:
            labeled_csv_path = config_file["evaluation"]["labeled_data_path"]
        else:
            print("No labeled data path specified. Please provide labeled_csv_path or update the config.")
            return None

    if not Path(pred_csv_path).exists():
        print(f"Prediction CSV file does not exist: {pred_csv_path}")
        return None

    if not Path(labeled_csv_path).exists():
        print(f"Labeled CSV file does not exist: {labeled_csv_path}")
        return None

    print("\n" + "=" * 60)
    print("Running record-based evaluation...")
    print("=" * 60)

    if output_dir is None:
        output_dir = config_file.get("results", {}).get("current_run_dir", None)

    record_based_results = evaluate_record_based_accuracy(
        pred_csv_path=pred_csv_path,
        labeled_csv_path=labeled_csv_path,
        config=config_file,
        lenient_mode=lenient_mode,
        output_dir=output_dir,
        ternary_complex_level=ternary_complex_level
    )

    return record_based_results


COMPOUND_MATCH_EXACT = "exact"
COMPOUND_MATCH_AMBIGUOUS = "ambiguous"
COMPOUND_MATCH_NO = "no"


def _match_compounds(row_a, row_b) -> str:
    """Match two rows by compound identity, returning a tri-state verdict"""
    a_name = row_a.get('Compound_Name')
    b_name = row_b.get('Compound_Name')
    a_ck = row_a.get('Connectivity_Key')
    b_ck = row_b.get('Connectivity_Key')

    a_name_missing = is_missing_value(a_name)
    b_name_missing = is_missing_value(b_name)
    a_ck_missing = is_missing_value(a_ck)
    b_ck_missing = is_missing_value(b_ck)

    both_ck_present = not a_ck_missing and not b_ck_missing
    ck_equal = both_ck_present and (
        normalize_dashes(str(a_ck).strip().lower())
        == normalize_dashes(str(b_ck).strip().lower())
    )

    if not a_name_missing and not b_name_missing:
        if normalize_dashes(str(a_name).strip().lower()) == normalize_dashes(str(b_name).strip().lower()):
            return COMPOUND_MATCH_EXACT
        if both_ck_present and ck_equal:
            return COMPOUND_MATCH_EXACT
        return COMPOUND_MATCH_NO

    if both_ck_present:
        return COMPOUND_MATCH_EXACT if ck_equal else COMPOUND_MATCH_NO

    return COMPOUND_MATCH_AMBIGUOUS


def _is_complementary_compound_ambiguity(row_a, row_b) -> bool:
    """One side has Compound_Name only (no Connectivity_Key), the other has"""
    a_name = not is_missing_value(row_a.get('Compound_Name'))
    b_name = not is_missing_value(row_b.get('Compound_Name'))
    a_ck = not is_missing_value(row_a.get('Connectivity_Key'))
    b_ck = not is_missing_value(row_b.get('Connectivity_Key'))
    return (a_name and not a_ck and not b_name and b_ck) or \
           (not a_name and a_ck and b_name and not b_ck)


def _identity_via_assay_fingerprint(row_a, row_b) -> bool:
    """Confirm record identity via DC50 (preferred) or Dmax fingerprint"""
    dc50_a = row_a.get('DC50')
    dc50_b = row_b.get('DC50')
    if not is_missing_value(dc50_a) and not is_missing_value(dc50_b):
        return compare_assay_values(dc50_a, dc50_b)

    dmax_a = row_a.get('Dmax')
    dmax_b = row_b.get('Dmax')
    if not is_missing_value(dmax_a) and not is_missing_value(dmax_b):
        return compare_assay_values(dmax_a, dmax_b)

    return False


def _no_field_contradictions(row_a, row_b, fields, config, cache_file_path, doi) -> bool:
    """True if no listed field contradicts"""
    for field in fields:
        val_a = row_a.get(field)
        val_b = row_b.get(field)
        if is_missing_value(val_a) or is_missing_value(val_b):
            continue
        is_match = check_semantic_match(
            pred_value=str(val_a).strip(),
            labeled_value=str(val_b).strip(),
            feature=field,
            config=config,
            cache_file_path=cache_file_path,
            doi=doi
        )
        if not is_match:
            return False
    return True


def _all_fields_positively_match(row_a, row_b, fields, config, cache_file_path, doi) -> bool:
    """All listed fields must be present on both sides AND match semantically"""
    for field in fields:
        val_a = row_a.get(field)
        val_b = row_b.get(field)
        if is_missing_value(val_a) or is_missing_value(val_b):
            return False
        is_match = check_semantic_match(
            pred_value=str(val_a).strip(),
            labeled_value=str(val_b).strip(),
            feature=field,
            config=config,
            cache_file_path=cache_file_path,
            doi=doi
        )
        if not is_match:
            return False
    return True


def _match_fields_semantic(row_a, row_b, fields, config, cache_file_path, doi) -> bool:
    """Check if all specified fields match between two rows using semantic matching"""
    for field in fields:
        val_a = row_a.get(field)
        val_b = row_b.get(field)

        a_missing = is_missing_value(val_a)
        b_missing = is_missing_value(val_b)

        if a_missing and b_missing:
            continue
        if a_missing or b_missing:
            return False

        is_match = check_semantic_match(
            pred_value=str(val_a).strip(),
            labeled_value=str(val_b).strip(),
            feature=field,
            config=config,
            cache_file_path=cache_file_path,
            doi=doi
        )
        if not is_match:
            return False

    return True


def _cascade_match_records(doi, baseline_df, llm_df, config, cache_file_path) -> dict:
    """Perform 3-level cascade matching between baseline and LLM records for a..."""
    L1_FIELDS = ['Degradation_Target', 'Recruiter', 'Assay', 'Cell_Line']
    L2_FIELDS = ['Degradation_Target', 'Recruiter']
    L3_FIELDS = ['Degradation_Target']

    baseline_records = [row.to_dict() for _, row in baseline_df.iterrows()]
    llm_pool = [row.to_dict() for _, row in llm_df.iterrows()]

    bl_active = dict(enumerate(baseline_records))
    llm_active = dict(enumerate(llm_pool))

    intersection_l1 = []
    intersection_l2 = []
    intersection_l3 = []

    def _greedy_assign(candidates):
        """Greedy non-conflicting assignment over (bl_idx, llm_idx, score, bl_rec, llm_rec)"""
        candidates_sorted = sorted(
            candidates, key=lambda c: (-c[2], c[0], c[1])
        )
        used_bl, used_llm = set(), set()
        chosen = []
        for bl_idx, llm_idx, _score, bl_rec, llm_rec in candidates_sorted:
            if bl_idx in used_bl or llm_idx in used_llm:
                continue
            used_bl.add(bl_idx)
            used_llm.add(llm_idx)
            chosen.append((bl_idx, llm_idx, bl_rec, llm_rec))
        return chosen

    def _resolve_level(predicate, bucket):
        """Collect every active (bl, llm) pair satisfying `predicate`, run a"""
        candidates = []
        for bl_idx, bl_rec in bl_active.items():
            for llm_idx, llm_rec in llm_active.items():
                if predicate(bl_rec, llm_rec):
                    score = calculate_assay_match_score(bl_rec, llm_rec)
                    candidates.append((bl_idx, llm_idx, score, bl_rec, llm_rec))
        for bl_idx, llm_idx, bl_rec, llm_rec in _greedy_assign(candidates):
            bucket.append({'baseline': bl_rec, 'llm': llm_rec})
            del bl_active[bl_idx]
            del llm_active[llm_idx]

    _resolve_level(
        lambda bl, lm: (
            _match_compounds(bl, lm) != COMPOUND_MATCH_NO
            and _match_fields_semantic(bl, lm, L1_FIELDS, config, cache_file_path, doi)
        ),
        intersection_l1,
    )

    _resolve_level(
        lambda bl, lm: (
            _match_compounds(bl, lm) == COMPOUND_MATCH_EXACT
            and _match_fields_semantic(bl, lm, L2_FIELDS, config, cache_file_path, doi)
        ),
        intersection_l2,
    )

    _resolve_level(
        lambda bl, lm: (
            _match_compounds(bl, lm) == COMPOUND_MATCH_EXACT
            and _match_fields_semantic(bl, lm, L3_FIELDS, config, cache_file_path, doi)
        ),
        intersection_l3,
    )

    fallback_buckets = {1: [], 2: [], 3: []}
    for bl_idx, bl_rec in bl_active.items():
        for llm_idx, llm_rec in llm_active.items():
            if not _is_complementary_compound_ambiguity(bl_rec, llm_rec):
                continue
            if not _identity_via_assay_fingerprint(bl_rec, llm_rec):
                continue
            if not _no_field_contradictions(bl_rec, llm_rec, L1_FIELDS, config, cache_file_path, doi):
                continue
            if not _all_fields_positively_match(bl_rec, llm_rec, L3_FIELDS, config, cache_file_path, doi):
                continue
            if _all_fields_positively_match(bl_rec, llm_rec, L1_FIELDS, config, cache_file_path, doi):
                level = 1
            elif _all_fields_positively_match(bl_rec, llm_rec, L2_FIELDS, config, cache_file_path, doi):
                level = 2
            else:
                level = 3
            score = calculate_assay_match_score(bl_rec, llm_rec)
            fallback_buckets[level].append((bl_idx, llm_idx, score, bl_rec, llm_rec))

    level_to_bucket = {1: intersection_l1, 2: intersection_l2, 3: intersection_l3}
    for level in (1, 2, 3):
        active_cands = [
            c for c in fallback_buckets[level]
            if c[0] in bl_active and c[1] in llm_active
        ]
        for bl_idx, llm_idx, bl_rec, llm_rec in _greedy_assign(active_cands):
            level_to_bucket[level].append({'baseline': bl_rec, 'llm': llm_rec})
            del bl_active[bl_idx]
            del llm_active[llm_idx]

    baseline_only = [{'baseline': bl_rec} for bl_rec in bl_active.values()]
    llm_only = [{'llm': rec} for rec in llm_active.values()]

    return {
        'intersection_l1': intersection_l1,
        'intersection_l2': intersection_l2,
        'intersection_l3': intersection_l3,
        'baseline_only': baseline_only,
        'llm_only': llm_only,
    }


def run_llm_vs_baseline_evaluation(config_file: dict,
                                    llm_csv_path: str = None,
                                    baseline_csv_path: str = None,
                                    output_dir: str = None):
    """Compare LLM extraction results against a baseline CSV using a 3-level"""
    print("-" * 60)
    print("Starting LLM vs Baseline evaluation (3-level cascade)...")

    if not llm_csv_path:
        json_file_path = Path(config_file["results"]["results_dir"])
        all_subdirs = [
            d for d in Path(json_file_path).iterdir()
            if d.is_dir() and len(d.name) == 11
            and d.name[:6].isdigit() and d.name[7:11].isdigit()
        ]
        if not all_subdirs:
            print("No result directory found. Please run the model extraction first.")
            return None

        latest_json_dir = sorted(all_subdirs, key=lambda d: d.name, reverse=True)[0]
        processed_csv_files = list(latest_json_dir.glob("*_processed.csv"))
        if not processed_csv_files:
            print(f"No processed CSV files found in {latest_json_dir}.")
            return None

        llm_csv_path = str(processed_csv_files[0])
        print(f"Using the latest processed CSV as LLM output: {llm_csv_path}")

    if not baseline_csv_path:
        if "evaluation" in config_file and "baseline_data_path" in config_file["evaluation"]:
            baseline_csv_path = config_file["evaluation"]["baseline_data_path"]
        else:
            print("No baseline data path specified. Please provide --baseline or set evaluation.baseline_data_path in config.")
            return None

    if not Path(llm_csv_path).exists():
        print(f"LLM CSV file does not exist: {llm_csv_path}")
        return None

    if not Path(baseline_csv_path).exists():
        print(f"Baseline CSV file does not exist: {baseline_csv_path}")
        return None

    print(f"LLM CSV (prediction): {llm_csv_path}")
    print(f"Baseline CSV (reference): {baseline_csv_path}")

    if output_dir is None:
        output_dir = config_file.get("results", {}).get("current_run_dir", None)

    cache_file_path = None
    if config_file and "evaluation" in config_file and "cache_file_path" in config_file["evaluation"]:
        config_cache_path = config_file["evaluation"]["cache_file_path"]
        if config_cache_path and str(config_cache_path).strip():
            cache_file_path = config_cache_path
    if cache_file_path is None and output_dir:
        cache_file_path = str(Path(output_dir) / "semantic_match_cache.json")
    elif cache_file_path is None:
        cache_file_path = str(Path(llm_csv_path).parent / "semantic_match_cache.json")
    Path(cache_file_path).parent.mkdir(parents=True, exist_ok=True)
    print(f"Semantic match cache: {cache_file_path}")

    llm_df = pd.read_csv(llm_csv_path)
    baseline_df = pd.read_csv(baseline_csv_path)
    llm_df.columns = llm_df.columns.str.strip()
    baseline_df.columns = baseline_df.columns.str.strip()

    print(f"\nTotal LLM records: {len(llm_df)}")
    print(f"Total Baseline records: {len(baseline_df)}")

    all_dois = set(baseline_df['DOI'].unique()) | set(llm_df['DOI'].unique())
    print(f"Total unique DOIs: {len(all_dois)}")

    print("\n" + "=" * 60)
    print("Running 3-level cascade matching...")
    print("=" * 60)

    global_counts = {'intersection_l1': 0, 'intersection_l2': 0, 'intersection_l3': 0,
                     'baseline_only': 0, 'llm_only': 0}
    per_doi_results = {}

    for doi in sorted(all_dois):
        doi_baseline = baseline_df[baseline_df['DOI'] == doi]
        doi_llm = llm_df[llm_df['DOI'] == doi]

        result = _cascade_match_records(doi, doi_baseline, doi_llm, config_file, cache_file_path)

        counts = {k: len(v) for k, v in result.items()}
        for k in global_counts:
            global_counts[k] += counts[k]

        per_doi_results[doi] = {
            'counts': counts,
            'records': result,
        }

        print(f"\n  DOI: {doi}")
        print(f"    Baseline records: {len(doi_baseline)}, LLM records: {len(doi_llm)}")
        print(f"    L1 (full):      {counts['intersection_l1']}")
        print(f"    L2 (mechanism): {counts['intersection_l2']}")
        print(f"    L3 (phenotype): {counts['intersection_l3']}")
        print(f"    Baseline only:  {counts['baseline_only']}")
        print(f"    LLM only:       {counts['llm_only']}")

    total_baseline = len(baseline_df)
    total_llm = len(llm_df)
    total_matched = global_counts['intersection_l1'] + global_counts['intersection_l2'] + global_counts['intersection_l3']

    print("\n" + "=" * 60)
    print("GLOBAL SUMMARY")
    print("=" * 60)
    print(f"  Intersection L1 (full match):      {global_counts['intersection_l1']}")
    print(f"  Intersection L2 (mechanism level):  {global_counts['intersection_l2']}")
    print(f"  Intersection L3 (phenotype level):  {global_counts['intersection_l3']}")
    print(f"  Total matched:                      {total_matched}")
    print(f"  Baseline only:                      {global_counts['baseline_only']}")
    print(f"  LLM only:                           {global_counts['llm_only']}")
    print(f"\n  Sanity check:")
    print(f"    Matched + Baseline_Only = {total_matched + global_counts['baseline_only']}  (should be {total_baseline})")
    print(f"    Matched + LLM_Only      = {total_matched + global_counts['llm_only']}  (should be {total_llm})")

    def _serialize_records(record_list):
        """Convert record dicts (may contain numpy/pandas types) to JSON-serializable..."""
        serialized = []
        for entry in record_list:
            s_entry = {}
            for side_key, side_val in entry.items():
                if isinstance(side_val, dict):
                    s_entry[side_key] = {
                        k: (None if pd.isna(v) else v) if not isinstance(v, str) else v
                        for k, v in side_val.items()
                    }
                else:
                    s_entry[side_key] = side_val
            serialized.append(s_entry)
        return serialized

    output_json = {
        'global_counts': global_counts,
        'total_baseline_records': total_baseline,
        'total_llm_records': total_llm,
        'per_doi': {}
    }

    for doi, doi_data in per_doi_results.items():
        output_json['per_doi'][doi] = {
            'counts': doi_data['counts'],
            'records': {
                k: _serialize_records(v) for k, v in doi_data['records'].items()
            }
        }

    if output_dir:
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        json_path = Path(output_dir) / "llm_vs_baseline_results.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(output_json, f, indent=2, ensure_ascii=False, default=str)
        print(f"\nResults saved to: {json_path}")

    return output_json


def query_chembl_properties(compound_name: str) -> Tuple[Dict[str, Any], str]:
    """Query all properties of a compound from ChEMBL"""
    try:
        from chembl_webresource_client.new_client import new_client
        
        molecule = new_client.molecule
        
        try:
            results = molecule.search(compound_name)
            
            if results:
                for result in results:
                    if isinstance(result, dict):
                        structures = result.get('molecule_structures', {}) or {}
                        properties = result.get('molecule_properties', {}) or {}
                        
                        if structures or properties:
                            props = {
                                'canonical_smiles': structures.get('canonical_smiles'),
                                'inchi': structures.get('standard_inchi'),
                                'inchikey': structures.get('standard_inchi_key'),
                                'molecular_formula': properties.get('full_molformula'),
                                'molecular_weight': properties.get('full_mwt'),
                                'alogp': properties.get('alogp'),
                                'psa': properties.get('psa'),
                                'hba': properties.get('hba'),
                                'hbd': properties.get('hbd'),
                                'rtb': properties.get('rtb'),
                                'num_ro5_violations': properties.get('num_ro5_violations'),
                                'aromatic_rings': properties.get('aromatic_rings'),
                                'heavy_atoms': properties.get('heavy_atoms'),
                                'iupac_name': result.get('pref_name'),
                            }
                            
                            props = {k: v for k, v in props.items() if v is not None}
                            
                            if props:
                                return props, "success"
        except Exception as search_error:
            print(f" [search failed: {str(search_error)[:30]}]", end="")
        
        try:
            results = molecule.filter(pref_name__iexact=compound_name)
            
            if results:
                for result in results:
                    if isinstance(result, dict):
                        structures = result.get('molecule_structures', {}) or {}
                        properties = result.get('molecule_properties', {}) or {}
                        
                        if structures or properties:
                            props = {
                                'canonical_smiles': structures.get('canonical_smiles'),
                                'inchi': structures.get('standard_inchi'),
                                'inchikey': structures.get('standard_inchi_key'),
                                'molecular_formula': properties.get('full_molformula'),
                                'molecular_weight': properties.get('full_mwt'),
                                'alogp': properties.get('alogp'),
                                'psa': properties.get('psa'),
                                'hba': properties.get('hba'),
                                'hbd': properties.get('hbd'),
                                'rtb': properties.get('rtb'),
                                'num_ro5_violations': properties.get('num_ro5_violations'),
                                'aromatic_rings': properties.get('aromatic_rings'),
                                'heavy_atoms': properties.get('heavy_atoms'),
                                'iupac_name': result.get('pref_name'),
                            }
                            
                            props = {k: v for k, v in props.items() if v is not None}
                            
                            if props:
                                return props, "success"
        except Exception as filter_error:
            print(f" [filter failed: {str(filter_error)[:30]}]", end="")
        
        return None, "not_found"
    
    except ImportError:
        print(" [ChEMBL client not installed]", end="")
        return None, "api_error"
    except Exception as e:
        print(f" [ChEMBL error: {str(e)[:50]}]", end="")
        return None, "api_error"



def fetch_smiles_multi_source(csv_file_path: str, 
                               use_chembl: bool = True,
                               rate_limit_pubchem: float = 0.2,
                               rate_limit_chembl: float = 0.5) -> Path:
    """Fetch SMILES from multiple sources (PubChem + ChEMBL)"""
    print("\n" + "="*60)
    print("SMILES EXTRACTION FROM MULTIPLE SOURCES")
    print("="*60)
    
    csv_path = Path(csv_file_path)
    output_path = csv_path.parent / "data_with_smiles.csv"
    failed_log_path_pubchem = csv_path.parent / "failed_smiles_lookup_pubchem.json"
    
    print(f"Reading CSV from: {csv_path}")
    df = pd.read_csv(csv_path)
    print(f"Total compounds: {len(df)}")
    
    property_columns = {
        'Canonical_SMILES': None,
        'SMILES_Source': None,
        'SMILES_Status': None,
        'InChI': None,
        'InChIKey': None,
        'Molecular_Formula': None,
        'Molecular_Weight': None,
        'XLogP': None,
        'TPSA': None,
        'Complexity': None,
        'HBond_Donor_Count': None,
        'HBond_Acceptor_Count': None,
        'Rotatable_Bond_Count': None,
        'Exact_Mass': None,
        'Monoisotopic_Mass': None,
        'Properties_Source': None
    }

    for col, default_val in property_columns.items():
        if col not in df.columns:
            df[col] = default_val
    
    failed_records_pubchem = []
    
    print("\n" + "="*60)
    print("PHASE 1: Querying PubChem")
    print("="*60)
    print(f"Rate limit: {rate_limit_pubchem}s between requests\n")
    
    for idx, row in df.iterrows():
        compound_name = row.get('Compound_Name', '')
        iupac_name = row.get('IUPAC_Name', '')
        
        query_name = None
        source = None
        
        if pd.notna(iupac_name) and str(iupac_name).strip():
            query_name = str(iupac_name).strip()
            source = 'IUPAC_Name'
        elif pd.notna(compound_name) and str(compound_name).strip():
            compound_str = str(compound_name).strip()
            if not compound_str.replace('.', '').replace('-', '').isdigit():
                query_name = compound_str
                source = 'Compound_Name'
        
        if not query_name:
            df.at[idx, 'SMILES_Status'] = 'skipped'
            print(f"  Row {idx}: Skipped (numeric or empty compound name)")
            continue
        
        try:
            query_display = query_name if len(query_name) <= 50 else query_name[:47] + "..."
            print(f"  Row {idx}: PubChem '{query_display}' ({source})...", end=" ")
            
            compounds = pcp.get_compounds(query_name, 'name')
            
            if compounds and len(compounds) > 0:
                compound = compounds[0]

                df.at[idx, 'Canonical_SMILES'] = getattr(compound, 'canonical_smiles', None)
                df.at[idx, 'InChI'] = getattr(compound, 'inchi', None)
                df.at[idx, 'InChIKey'] = getattr(compound, 'inchikey', None)
                df.at[idx, 'Molecular_Formula'] = getattr(compound, 'molecular_formula', None)
                df.at[idx, 'Molecular_Weight'] = getattr(compound, 'molecular_weight', None)
                df.at[idx, 'XLogP'] = getattr(compound, 'xlogp', None)
                df.at[idx, 'TPSA'] = getattr(compound, 'tpsa', None)
                df.at[idx, 'Complexity'] = getattr(compound, 'complexity', None)
                df.at[idx, 'HBond_Donor_Count'] = getattr(compound, 'h_bond_donor_count', None)
                df.at[idx, 'HBond_Acceptor_Count'] = getattr(compound, 'h_bond_acceptor_count', None)
                df.at[idx, 'Rotatable_Bond_Count'] = getattr(compound, 'rotatable_bond_count', None)
                df.at[idx, 'Exact_Mass'] = getattr(compound, 'exact_mass', None)
                df.at[idx, 'Monoisotopic_Mass'] = getattr(compound, 'monoisotopic_mass', None)

                if not pd.notna(df.at[idx, 'IUPAC_Name']) or not str(df.at[idx, 'IUPAC_Name']).strip():
                    pubchem_iupac = getattr(compound, 'iupac_name', None)
                    if pubchem_iupac:
                        df.at[idx, 'IUPAC_Name'] = pubchem_iupac

                df.at[idx, 'SMILES_Source'] = source
                df.at[idx, 'SMILES_Status'] = 'success'
                df.at[idx, 'Properties_Source'] = 'pubchem'

                smiles = df.at[idx, 'Canonical_SMILES']
                smiles_display = smiles[:60] if smiles else "N/A"
                print(f"✓ {smiles_display}")
            else:
                df.at[idx, 'SMILES_Status'] = 'not_found'
                failed_records_pubchem.append({
                    'row': idx,
                    'DOI': row.get('DOI', ''),
                    'Compound_Name': compound_name,
                    'IUPAC_Name': iupac_name,
                    'query_name': query_name,
                    'source': source,
                    'status': 'not_found',
                    'error': 'No compounds found in PubChem'
                })
                print("✗ Not found")
        
        except pcp.NotFoundError:
            df.at[idx, 'SMILES_Status'] = 'not_found'
            failed_records_pubchem.append({
                'row': idx,
                'DOI': row.get('DOI', ''),
                'Compound_Name': compound_name,
                'IUPAC_Name': iupac_name,
                'query_name': query_name,
                'source': source,
                'status': 'not_found',
                'error': 'PubChem NotFoundError'
            })
            print("✗ Not found in PubChem")
        
        except Exception as e:
            df.at[idx, 'SMILES_Status'] = 'api_error'
            failed_records_pubchem.append({
                'row': idx,
                'DOI': row.get('DOI', ''),
                'Compound_Name': compound_name,
                'IUPAC_Name': iupac_name,
                'query_name': query_name,
                'source': source,
                'status': 'error',
                'error': str(e)
            })
            print(f"✗ Error: {str(e)[:50]}")
        
        time.sleep(rate_limit_pubchem)
    
    df.to_csv(output_path, index=False)
    
    if failed_records_pubchem:
        with open(failed_log_path_pubchem, 'w', encoding='utf-8') as f:
            json.dump(failed_records_pubchem, f, ensure_ascii=False, indent=2)
    
    pubchem_success = len(df[df['SMILES_Status'] == 'success'])
    pubchem_failed = len(failed_records_pubchem)
    print(f"\nPhase 1 Results: {pubchem_success} success, {pubchem_failed} failed")
    
    if not use_chembl:
        print("\nSkipping ChEMBL phase (--no-chembl flag used)")
        print(f"\n✓ Saved SMILES data to: {output_path}")
        return output_path
    
    print("\n" + "="*60)
    print("PHASE 2: Querying ChEMBL for Failed Compounds")
    print("="*60)
    
    failed_rows = df[df['SMILES_Status'] == 'not_found'].index.tolist()
    print(f"Found {len(failed_rows)} compounds to retry with ChEMBL")
    print(f"Rate limit: {rate_limit_chembl}s between requests\n")
    
    if len(failed_rows) == 0:
        print("No compounds to query in ChEMBL\n")
        print(f"✓ Saved SMILES data to: {output_path}")
        return output_path
    
    chembl_success = 0
    chembl_failed = []
    
    for idx in failed_rows:
        row = df.loc[idx]
        compound_name = row.get('Compound_Name', '')
        iupac_name = row.get('IUPAC_Name', '')
        
        query_name = None
        source = None
        
        if pd.notna(iupac_name) and str(iupac_name).strip():
            query_name = str(iupac_name).strip()
            source = 'IUPAC_Name'
        elif pd.notna(compound_name) and str(compound_name).strip():
            compound_str = str(compound_name).strip()
            if not compound_str.replace('.', '').replace('-', '').isdigit():
                query_name = compound_str
                source = 'Compound_Name'
        
        if not query_name:
            continue
        
        query_display = query_name if len(query_name) <= 50 else query_name[:47] + "..."
        print(f"  Row {idx}: ChEMBL '{query_display}'...", end="")
        
        props, status = query_chembl_properties(query_name)

        if status == "success" and props:
            if not pd.notna(df.at[idx, 'Canonical_SMILES']) and props.get('canonical_smiles'):
                df.at[idx, 'Canonical_SMILES'] = props.get('canonical_smiles')

            if not pd.notna(df.at[idx, 'InChI']) and props.get('inchi'):
                df.at[idx, 'InChI'] = props.get('inchi')

            if not pd.notna(df.at[idx, 'InChIKey']) and props.get('inchikey'):
                df.at[idx, 'InChIKey'] = props.get('inchikey')

            if not pd.notna(df.at[idx, 'Molecular_Formula']) and props.get('molecular_formula'):
                df.at[idx, 'Molecular_Formula'] = props.get('molecular_formula')

            if not pd.notna(df.at[idx, 'Molecular_Weight']) and props.get('molecular_weight'):
                df.at[idx, 'Molecular_Weight'] = props.get('molecular_weight')

            if not pd.notna(df.at[idx, 'XLogP']) and props.get('alogp'):
                df.at[idx, 'XLogP'] = props.get('alogp')

            if not pd.notna(df.at[idx, 'TPSA']) and props.get('psa'):
                df.at[idx, 'TPSA'] = props.get('psa')

            if not pd.notna(df.at[idx, 'HBond_Acceptor_Count']) and props.get('hba'):
                df.at[idx, 'HBond_Acceptor_Count'] = props.get('hba')

            if not pd.notna(df.at[idx, 'HBond_Donor_Count']) and props.get('hbd'):
                df.at[idx, 'HBond_Donor_Count'] = props.get('hbd')

            if not pd.notna(df.at[idx, 'Rotatable_Bond_Count']) and props.get('rtb'):
                df.at[idx, 'Rotatable_Bond_Count'] = props.get('rtb')

            if (not pd.notna(df.at[idx, 'IUPAC_Name']) or not str(df.at[idx, 'IUPAC_Name']).strip()) and props.get('iupac_name'):
                df.at[idx, 'IUPAC_Name'] = props.get('iupac_name')

            df.at[idx, 'SMILES_Source'] = f"{source}_chembl"
            df.at[idx, 'SMILES_Status'] = 'success_chembl'
            df.at[idx, 'Properties_Source'] = 'chembl'
            chembl_success += 1

            smiles = props.get('canonical_smiles', 'N/A')
            smiles_display = smiles[:60] if len(smiles) > 60 else smiles
            print(f" ✓ {smiles_display}")
        else:
            chembl_failed.append({
                'row': idx,
                'DOI': row.get('DOI', ''),
                'Compound_Name': compound_name,
                'IUPAC_Name': iupac_name,
                'query_name': query_name,
                'source': source,
                'status': status,
                'error': 'ChEMBL query failed'
            })
            print(" ✗ Not found")
        
        time.sleep(rate_limit_chembl)
    
    df.to_csv(output_path, index=False)
    print(f"\n✓ Updated CSV with ChEMBL results: {output_path}")
    
    if chembl_failed:
        failed_log_path_final = csv_path.parent / "failed_smiles_lookup_after_chembl.json"
        with open(failed_log_path_final, 'w', encoding='utf-8') as f:
            json.dump(chembl_failed, f, ensure_ascii=False, indent=2)
        print(f"✓ Saved final failed lookup log: {failed_log_path_final}")
    
    print(f"\n{'-'*60}")
    print("FINAL SMILES EXTRACTION STATISTICS")
    print(f"{'-'*60}")
    print(f"  Phase 1 (PubChem):")
    print(f"    Success: {pubchem_success}")
    print(f"    Failed:  {pubchem_failed}")
    print(f"\n  Phase 2 (ChEMBL):")
    print(f"    Success: {chembl_success}")
    print(f"    Failed:  {len(chembl_failed)}")
    print(f"\n  Overall:")
    total_success = pubchem_success + chembl_success
    total_failed = len(chembl_failed)
    total_compounds = len(df[df['SMILES_Status'] != 'skipped'])
    success_rate = (total_success / total_compounds * 100) if total_compounds > 0 else 0
    print(f"    Total success: {total_success}/{total_compounds}")
    print(f"    Total failed:  {total_failed}")
    print(f"    Success rate:  {success_rate:.1f}%")
    print(f"{'-'*60}")
    
    return output_path


